from __future__ import annotations

import json
import math
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Tuple

import pandas as pd
import streamlit as st

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except Exception:  # pragma: no cover - l'export DOCX reste optionnel si la dépendance manque
    Document = None
    OxmlElement = None
    qn = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
except Exception:  # pragma: no cover - l'export Excel reste optionnel si la dépendance manque
    Workbook = None

st.set_page_config(
    page_title="Système expert CGP – Holding familiale",
    page_icon="🏛️",
    layout="wide",
)

UNKNOWN = "Non renseigné"
NA = "Non applicable"
YES = "Oui"
NO = "Non"
UNCERTAIN = "Incertain"

# =============================================================================
# 1. Données de référence : risques, conséquences, outils, actions
# =============================================================================

@dataclass
class RiskDefinition:
    code: str
    objectif: str
    libelle: str
    consequences: List[str]
    outils: List[str]
    actions_preventives: List[str]
    professionnels: List[str]
    gravite_base: int


RISK_DEFINITIONS: Dict[str, RiskDefinition] = {
    "dilution": RiskDefinition(
        "dilution", "Pérennité / contrôle", "Dilution progressive du capital",
        ["Dispersion des titres entre plusieurs branches familiales", "Affaiblissement du pouvoir de décision familial", "Risque de perte de cohérence stratégique"],
        ["Holding familiale", "Clause d’agrément", "Clause de préemption", "Actions de préférence", "Pacte d’associés"],
        ["Cartographier les futurs associés", "Encadrer les entrées et sorties du capital", "Anticiper les transmissions successives"],
        ["CGP", "Avocat", "Notaire"], 4,
    ),
    "tiers": RiskDefinition(
        "tiers", "Pérennité / contrôle", "Entrée de tiers au capital",
        ["Perte d’indépendance familiale", "Conflit stratégique avec un associé extérieur", "Remise en cause du projet familial"],
        ["Clause d’agrément", "Clause de préemption", "Pacte d’associés", "Statuts adaptés"],
        ["Formaliser les règles de cession", "Prévoir un droit de rachat prioritaire", "Vérifier la cohérence statuts / pacte"],
        ["CGP", "Avocat"], 4,
    ),
    "blocage_gouvernance": RiskDefinition(
        "blocage_gouvernance", "Pérennité / contrôle", "Blocage de gouvernance",
        ["Paralysie décisionnelle", "Conflits entre associés actifs et passifs", "Perte de valeur ou ralentissement de l’entreprise"],
        ["Pacte d’associés", "Charte familiale", "Conseil de famille", "Comité stratégique", "Statuts de SAS"],
        ["Clarifier les rôles", "Définir les règles de majorité", "Organiser un dialogue familial"],
        ["CGP", "Avocat", "Conseil en gouvernance familiale"], 4,
    ),
    "successeur": RiskDefinition(
        "successeur", "Pérennité / contrôle", "Absence ou mauvaise préparation du successeur",
        ["Rupture de continuité managériale", "Perte de confiance des salariés et partenaires", "Cession subie ou direction externe non anticipée"],
        ["Gouvernance transitoire", "Direction externe temporaire", "Actions de préférence", "Pacte d’associés", "Calendrier de transmission"],
        ["Identifier les intentions des héritiers", "Évaluer les compétences du successeur", "Prévoir un calendrier progressif"],
        ["CGP", "Expert-comptable", "Conseil en gouvernance", "Avocat"], 5,
    ),
    "conflit_heritiers": RiskDefinition(
        "conflit_heritiers", "Équité entre héritiers", "Conflit repreneur / non repreneurs",
        ["Sentiment d’injustice", "Blocage de la holding", "Tensions durables entre branches familiales"],
        ["Donation-partage", "Soulte", "Pacte d’associés", "Mécanisme de sortie", "Droits financiers différenciés"],
        ["Distinguer égalité et équité", "Organiser une valorisation indépendante", "Anticiper la liquidité des non repreneurs", "Expliquer les choix avant les actes"],
        ["CGP", "Notaire", "Avocat", "Expert-comptable"], 4,
    ),
    "contestation": RiskDefinition(
        "contestation", "Équité entre héritiers", "Contestation successorale",
        ["Remise en cause des opérations de transmission", "Contentieux familial", "Instabilité juridique de la holding"],
        ["Donation-partage", "Audit civil", "Valorisation indépendante", "Démembrement", "Testament"],
        ["Vérifier la réserve héréditaire", "Documenter la valorisation", "Faire intervenir le notaire dès le diagnostic"],
        ["CGP", "Notaire", "Avocat", "Expert-comptable"], 4,
    ),
    "liquidite": RiskDefinition(
        "liquidite", "Équité entre héritiers", "Soulte trop lourde ou manque de liquidité",
        ["Difficulté de paiement des soultes ou droits", "Pression excessive sur les dividendes", "Cession forcée d’actifs"],
        ["Soulte échelonnée", "Family Buy-Out", "Paiement différé/fractionné", "Assurance-vie", "Politique de distribution encadrée"],
        ["Simuler les besoins de liquidité", "Mesurer la capacité de financement", "Prévoir les conditions de sortie des héritiers non actifs"],
        ["CGP", "Expert-comptable", "Notaire", "Avocat"], 4,
    ),
    "fiscalite": RiskDefinition(
        "fiscalite", "Optimisation fiscale", "Charge fiscale excessive",
        ["Droits de mutation élevés", "Besoin de liquidité immédiate", "Risque de cession partielle pour financer la fiscalité"],
        ["Pacte Dutreil", "Donation anticipée", "Démembrement", "Paiement différé et fractionné"],
        ["Réaliser une simulation fiscale", "Anticiper la transmission", "Évaluer la fiscalité résiduelle"],
        ["CGP", "Notaire", "Avocat fiscaliste", "Expert-comptable"], 4,
    ),
    "dutreil": RiskDefinition(
        "dutreil", "Optimisation fiscale", "Remise en cause du Pacte Dutreil",
        ["Perte de l’exonération partielle", "Imposition complémentaire imprévue", "Fragilisation financière des héritiers"],
        ["Audit Dutreil", "Documentation holding animatrice", "Suivi des engagements", "Revue juridique annuelle"],
        ["Vérifier l’éligibilité", "Documenter la qualité de holding animatrice", "Suivre les engagements de conservation"],
        ["CGP", "Avocat fiscaliste", "Notaire", "Expert-comptable"], 5,
    ),
    "conjoint": RiskDefinition(
        "conjoint", "Protection familiale", "Fragilisation du conjoint et des proches",
        ["Insuffisance de revenus du conjoint survivant", "Dépendance aux décisions des enfants", "Tensions conjoint / héritiers"],
        ["Régime matrimonial", "Donation entre époux", "Assurance-vie", "Démembrement", "Prévoyance"],
        ["Analyser les droits du conjoint", "Simuler un décès ou une incapacité", "Organiser des revenus ou capitaux disponibles"],
        ["CGP", "Notaire", "Assureur", "Avocat"], 4,
    ),
    "dependance": RiskDefinition(
        "dependance", "Protection familiale", "Dépendance excessive au patrimoine professionnel",
        ["Exposition aux aléas économiques de l’entreprise", "Manque de liquidité familiale", "Difficulté à protéger conjoint et héritiers"],
        ["Diversification patrimoniale", "Holding patrimoniale", "Assurance-vie", "Prévoyance", "Investissements complémentaires"],
        ["Mesurer le poids de l’entreprise", "Constituer des actifs liquides", "Prévoir des scénarios de crise"],
        ["CGP", "Expert-comptable", "Notaire"], 4,
    ),
    "suivi": RiskDefinition(
        "suivi", "Suivi", "Inadaptation progressive de la stratégie",
        ["Outils devenus inadaptés", "Non-respect d’engagements fiscaux ou juridiques", "Perte de cohérence dans le temps"],
        ["Rendez-vous de suivi", "Revue des statuts", "Actualisation du pacte", "Revue patrimoniale annuelle"],
        ["Planifier une revue annuelle", "Réévaluer les objectifs familiaux", "Contrôler les engagements"],
        ["CGP", "Notaire", "Avocat", "Expert-comptable"], 3,
    ),
}


TOOL_JUSTIFICATIONS_BY_RISK: Dict[str, Dict[str, str]] = {
    "dilution": {
        "Holding familiale": "Elle centralise la détention des titres et limite la dispersion directe du capital entre plusieurs branches familiales.",
        "Clause d’agrément": "Elle permet de contrôler l’entrée de nouveaux associés et d’éviter qu’un tiers non souhaité intègre le capital.",
        "Clause de préemption": "Elle donne une priorité de rachat aux associés familiaux avant toute cession à un tiers.",
        "Actions de préférence": "Elles peuvent renforcer certains droits politiques afin de préserver le contrôle familial malgré l’évolution de l’actionnariat.",
        "Pacte d’associés": "Il formalise les règles d’entrée, de sortie et de conservation des titres entre membres de la famille.",
    },
    "tiers": {
        "Clause d’agrément": "Elle soumet l’entrée d’un nouvel associé à l’accord préalable des associés ou de l’organe compétent.",
        "Clause de préemption": "Elle permet aux associés familiaux de se porter acquéreurs en priorité avant une cession extérieure.",
        "Pacte d’associés": "Il encadre contractuellement les cessions et réduit le risque d’entrée non anticipée d’un tiers.",
        "Statuts adaptés": "Ils permettent d’intégrer directement les règles de contrôle des cessions dans l’organisation juridique de la société.",
    },
    "blocage_gouvernance": {
        "Pacte d’associés": "Il permet d’organiser les règles de décision, d’information, de sortie et de résolution des désaccords.",
        "Charte familiale": "Elle clarifie les valeurs, les rôles et les principes de fonctionnement de la famille autour de l’entreprise.",
        "Conseil de famille": "Il crée un espace de dialogue distinct de la gestion opérationnelle et limite la transformation des désaccords en blocages.",
        "Comité stratégique": "Il permet d’associer certaines parties prenantes aux grandes orientations sans confondre gouvernance familiale et direction opérationnelle.",
        "Statuts de SAS": "La souplesse statutaire de la SAS permet d’adapter les organes de décision aux besoins de la famille.",
    },
    "successeur": {
        "Gouvernance transitoire": "Elle accompagne le passage progressif du pouvoir entre le dirigeant et le repreneur.",
        "Direction externe temporaire": "Elle constitue une solution de continuité lorsque le successeur familial n’est pas encore prêt ou identifié.",
        "Actions de préférence": "Elles peuvent attribuer des droits adaptés au repreneur ou organiser une montée en puissance progressive.",
        "Pacte d’associés": "Il encadre la période de transition, les droits du repreneur et les équilibres avec les autres héritiers.",
        "Calendrier de transmission": "Il rend la transmission opérationnelle en fixant les étapes de transfert du pouvoir et des responsabilités.",
    },
    "conflit_heritiers": {
        "Donation-partage": "Elle organise la répartition du patrimoine du vivant du dirigeant et réduit le risque de contestation entre héritiers.",
        "Soulte": "Elle permet de compenser financièrement les héritiers non repreneurs lorsque l’entreprise est attribuée principalement au repreneur.",
        "Pacte d’associés": "Il fixe les droits et obligations des héritiers associés, notamment les règles de sortie, d’information et de distribution.",
        "Mécanisme de sortie": "Il évite que les héritiers non repreneurs se sentent enfermés dans une structure peu liquide.",
        "Droits financiers différenciés": "Ils peuvent concilier pouvoir renforcé du repreneur et protection économique des héritiers non actifs.",
    },
    "contestation": {
        "Donation-partage": "Elle permet d’anticiper le partage et de stabiliser les équilibres successoraux avant l’ouverture de la succession.",
        "Audit civil": "Il vérifie la cohérence de la stratégie avec la réserve héréditaire, les droits du conjoint et la situation familiale.",
        "Valorisation indépendante": "Elle réduit le risque de contestation sur la valeur des titres transmis.",
        "Démembrement": "Il permet d’organiser une transmission progressive tout en maintenant certains droits au profit du dirigeant ou du conjoint.",
        "Testament": "Il complète l’organisation successorale lorsque certaines volontés doivent être formalisées hors donation.",
    },
    "liquidite": {
        "Soulte échelonnée": "Elle évite une charge immédiate excessive pour le repreneur ou la holding.",
        "Family Buy-Out": "Il structure le financement de la reprise familiale et peut faciliter la compensation des héritiers non repreneurs.",
        "Paiement différé/fractionné": "Il étale le paiement de certaines charges fiscales et réduit la pression immédiate sur la trésorerie.",
        "Assurance-vie": "Elle peut fournir des capitaux disponibles pour compenser ou sécuriser certains héritiers.",
        "Politique de distribution encadrée": "Elle donne de la visibilité aux associés familiaux sans imposer une distribution désorganisée.",
    },
    "fiscalite": {
        "Pacte Dutreil": "Il peut réduire significativement l’assiette taxable de la transmission si les conditions sont respectées.",
        "Donation anticipée": "Elle permet d’organiser la transmission dans le temps et de mieux maîtriser la charge fiscale.",
        "Démembrement": "Il peut réduire la base taxable de la nue-propriété tout en maintenant certains droits économiques.",
        "Paiement différé et fractionné": "Il n’allège pas la fiscalité, mais en facilite le financement dans le temps.",
    },
    "dutreil": {
        "Audit Dutreil": "Il permet de vérifier l’éligibilité au régime avant la mise en œuvre et de réduire le risque de remise en cause.",
        "Documentation holding animatrice": "Elle sécurise la qualification de la holding lorsque le bénéfice du régime dépend de son rôle d’animation.",
        "Suivi des engagements": "Il garantit que les conditions de conservation et les obligations associées restent respectées dans le temps.",
        "Revue juridique annuelle": "Elle permet d’identifier les opérations ou évolutions susceptibles d’affecter la sécurité du dispositif.",
    },
    "conjoint": {
        "Régime matrimonial": "Son analyse permet d’identifier les droits du conjoint et d’adapter la protection du couple avant la transmission.",
        "Donation entre époux": "Elle peut augmenter les droits du conjoint survivant et offrir davantage de souplesse successorale.",
        "Assurance-vie": "Elle peut transmettre des capitaux disponibles au conjoint en dehors du règlement successoral ordinaire.",
        "Démembrement": "Il peut préserver des revenus ou une jouissance au profit du conjoint tout en préparant la transmission aux enfants.",
        "Prévoyance": "Elle apporte une liquidité rapide en cas de décès ou d’incapacité du dirigeant.",
    },
    "dependance": {
        "Diversification patrimoniale": "Elle réduit la dépendance de la famille à la seule valeur de l’entreprise.",
        "Holding patrimoniale": "Elle peut structurer progressivement des actifs distincts de l’entreprise opérationnelle.",
        "Assurance-vie": "Elle constitue un support de liquidité et de transmission complémentaire au patrimoine professionnel.",
        "Prévoyance": "Elle protège la famille contre un événement brutal affectant le dirigeant ou l’entreprise.",
        "Investissements complémentaires": "Ils permettent de constituer des sources de revenus et de liquidité en dehors de l’entreprise familiale.",
    },
    "suivi": {
        "Rendez-vous de suivi": "Ils permettent de vérifier régulièrement que la stratégie reste adaptée à la famille et à l’entreprise.",
        "Revue des statuts": "Elle évite que les règles de gouvernance deviennent inadaptées à l’évolution de l’actionnariat familial.",
        "Actualisation du pacte": "Elle permet d’ajuster les règles entre associés lorsque la situation familiale ou patrimoniale évolue.",
        "Revue patrimoniale annuelle": "Elle contrôle la cohérence globale entre objectifs, risques, outils et besoins de liquidité.",
    },
}

DEFAULT_ANSWERS: Dict[str, Any] = {
    "client_name": "",
    "client_age": 0,
    "company_name": "",
    "company_activity": "",
    "company_form": UNKNOWN,
    "cgp_name": "",
    "entretien_date": "",
    "rapport_orientation": "Équilibré",
    "niveau_detail": "Détaillé",
    "objectif_libre": "",
    "attentes_client": "",
    "contraintes_client": "",
    "personnes_a_associer": "",
    "observations": "",
    "maturite_projet": UNKNOWN,
    "delai_transmission": UNKNOWN,
    "urgence_evenement": UNKNOWN,
    "qualite_information": UNKNOWN,
    "objectifs": [],
    "objective_weights": {},
    "nb_enfants": 0,
    "conjoint_present": UNKNOWN,
    "famille_recomposee": UNKNOWN,
    "dialogue_familial": UNKNOWN,
    "accord_conjoint": UNKNOWN,
    "heritier_repreneur": UNKNOWN,
    "autres_heritiers_actifs": UNKNOWN,
    "volonte_non_repreneurs": UNKNOWN,
    "soulte_envisagee": UNKNOWN,
    "capacite_financement_soulte": UNKNOWN,
    "valorisation_independante": UNKNOWN,
    "audit_civil": UNKNOWN,
    "conjoint_dependant": UNKNOWN,
    "protection_conjoint_prevue": UNKNOWN,
    "regime_matrimonial_adapte": UNKNOWN,
    "valeur_entreprise": 0,
    "poids_entreprise": 0,
    "actifs_liquides": UNKNOWN,
    "endettement_familial": UNKNOWN,
    "besoin_revenus_famille": UNKNOWN,
    "diversification": UNKNOWN,
    "prevoyance": UNKNOWN,
    "gouvernance_formalisee": UNKNOWN,
    "associes_actifs_passifs": UNKNOWN,
    "clauses_entree_sortie": UNKNOWN,
    "politique_dividendes_definie": UNKNOWN,
    "entreprise_dependante_dirigeant": UNKNOWN,
    "calendrier_transmission": UNKNOWN,
    "successeur_prepare": UNKNOWN,
    "pacte_dutreil": UNKNOWN,
    "holding_animatrice": UNKNOWN,
    "simulation_fiscale": UNKNOWN,
    "audit_dutreil": UNKNOWN,
    "suivi_engagements": UNKNOWN,
    "suivi_annuel": UNKNOWN,
    "formalisation_rapport": UNKNOWN,
}

STEP_KEYS: Dict[int, List[str]] = {
    1: ["client_name", "client_age", "company_name", "company_activity", "company_form", "cgp_name", "entretien_date", "rapport_orientation", "niveau_detail", "objectif_libre", "attentes_client", "contraintes_client", "personnes_a_associer", "observations", "maturite_projet", "delai_transmission", "urgence_evenement", "qualite_information", "objectifs", "objective_weights"],
    2: ["nb_enfants", "conjoint_present", "famille_recomposee", "dialogue_familial", "accord_conjoint", "heritier_repreneur", "autres_heritiers_actifs", "volonte_non_repreneurs", "soulte_envisagee", "capacite_financement_soulte", "conjoint_dependant", "protection_conjoint_prevue", "regime_matrimonial_adapte", "valorisation_independante", "audit_civil"],
    3: ["valeur_entreprise", "poids_entreprise", "actifs_liquides", "endettement_familial", "besoin_revenus_famille", "diversification", "prevoyance"],
    4: ["clauses_entree_sortie", "gouvernance_formalisee", "associes_actifs_passifs", "politique_dividendes_definie", "entreprise_dependante_dirigeant", "successeur_prepare", "calendrier_transmission"],
    5: ["simulation_fiscale", "pacte_dutreil", "holding_animatrice", "audit_dutreil", "suivi_engagements", "suivi_annuel", "formalisation_rapport"],
    6: [],
}

PRIORITY_ORDER = {"Critique": 4, "Élevé": 3, "Moyen": 2, "Faible": 1, "Inexistant": 0}
PRIORITY_COLORS = {"Critique": "#7f1d1d", "Élevé": "#b45309", "Moyen": "#1d4ed8", "Faible": "#047857", "Inexistant": "#6b7280"}

OBJECTIVE_WEIGHT_OPTIONS = ["Important", "Très important", "Prioritaire"]
OBJECTIVE_WEIGHT_FACTORS = {"Important": 1.0, "Très important": 1.2, "Prioritaire": 1.4}
OBJECTIVE_DISPLAY_ORDER = [
    "Conserver le contrôle familial",
    "Transmettre l’entreprise",
    "Optimiser la fiscalité",
    "Protéger le conjoint et les proches",
    "Préserver l’équité entre les héritiers",
    "Diversifier le patrimoine",
]
RISK_TO_OBJECTIVE = {
    "dilution": "Conserver le contrôle familial",
    "tiers": "Conserver le contrôle familial",
    "blocage_gouvernance": "Conserver le contrôle familial",
    "successeur": "Transmettre l’entreprise",
    "conflit_heritiers": "Préserver l’équité entre les héritiers",
    "contestation": "Préserver l’équité entre les héritiers",
    "liquidite": "Préserver l’équité entre les héritiers",
    "fiscalite": "Optimiser la fiscalité",
    "dutreil": "Optimiser la fiscalité",
    "conjoint": "Protéger le conjoint et les proches",
    "dependance": "Diversifier le patrimoine",
    "suivi": "Suivi global",
}

# =============================================================================
# 2. Initialisation, widgets, validation explicite
# =============================================================================

def safe_key(text: str) -> str:
    return "".join(ch.lower() if ch.isalnum() else "_" for ch in text).strip("_")

def get_objective_weight(answers: Dict[str, Any], objective: str) -> str:
    weights = answers.get("objective_weights") or {}
    return weights.get(objective, "Important")

def get_objective_factor(answers: Dict[str, Any], objective: str) -> float:
    return OBJECTIVE_WEIGHT_FACTORS.get(get_objective_weight(answers, objective), 1.0)

def weight_for_risk(code: str, answers: Dict[str, Any] | None) -> Tuple[str, float]:
    answers = answers or {}
    objectives = list(answers.get("objectifs") or [])
    if code == "suivi":
        if not objectives:
            return "Important", 1.0
        labels = [get_objective_weight(answers, obj) for obj in objectives]
        label = max(labels, key=lambda lab: OBJECTIVE_WEIGHT_FACTORS.get(lab, 1.0))
        return f"Suivi global ({label.lower()})", OBJECTIVE_WEIGHT_FACTORS.get(label, 1.0)
    objective = RISK_TO_OBJECTIVE.get(code, "")
    if code == "liquidite" and "Diversifier le patrimoine" in objectives:
        objective = "Diversifier le patrimoine"
    if code == "dependance" and "Diversifier le patrimoine" not in objectives and "Protéger le conjoint et les proches" in objectives:
        objective = "Protéger le conjoint et les proches"
    label = get_objective_weight(answers, objective) if objective else "Important"
    factor = get_objective_factor(answers, objective) if objective else 1.0
    return label, factor

def init_app() -> None:
    if "answers" not in st.session_state:
        st.session_state.answers = dict(DEFAULT_ANSWERS)
    if "draft_answers" not in st.session_state:
        # Les réponses en cours de saisie sont conservées séparément des réponses validées.
        # Cela permet de revenir en arrière dans le questionnaire sans perdre les champs déjà remplis.
        st.session_state.draft_answers = dict(st.session_state.answers)
    if "validated_steps" not in st.session_state:
        st.session_state.validated_steps = set()
    if "last_validation" not in st.session_state:
        st.session_state.last_validation = {}
    st.session_state.setdefault("current_step", 1)
    st.session_state.setdefault("app_page", "Questionnaire adaptatif")
    for key, value in DEFAULT_ANSWERS.items():
        st.session_state.setdefault(f"w_{key}", st.session_state.draft_answers.get(key, st.session_state.answers.get(key, value)))


def reset_app() -> None:
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    init_app()




# =============================================================================
# 2. (suite) Sync, validation, widgets
# =============================================================================

def sync_draft_key(key: str) -> None:
    """Copie immédiatement la valeur d'un widget dans le brouillon."""
    widget_key = f"w_{key}"
    if widget_key in st.session_state:
        st.session_state.draft_answers[key] = st.session_state[widget_key]


def sync_widget_from_draft(key: str) -> None:
    """Initialise le widget avec la dernière valeur saisie."""
    value = st.session_state.draft_answers.get(
        key, st.session_state.answers.get(key, DEFAULT_ANSWERS.get(key))
    )
    st.session_state.setdefault(f"w_{key}", value)


def set_answer_and_draft(key: str, value: Any) -> None:
    st.session_state.answers[key] = value
    st.session_state.draft_answers[key] = value
    widget_key = f"w_{key}"
    if widget_key in st.session_state:
        st.session_state[widget_key] = value


def get_draft(key: str) -> Any:
    if f"w_{key}" in st.session_state:
        st.session_state.draft_answers[key] = st.session_state[f"w_{key}"]
    return st.session_state.draft_answers.get(
        key, st.session_state.answers.get(key, DEFAULT_ANSWERS.get(key))
    )


def selected_objectives_from_draft() -> set:
    return set(get_draft("objectifs") or [])


def save_step(step: int) -> Tuple[bool, str]:
    """Copie les réponses du brouillon vers answers. Seules answers sont utilisées pour le scoring."""
    keys = STEP_KEYS.get(step, [])
    if step == 1:
        # Restaure les objectifs depuis cache/answers si draft les a perdus
        if not st.session_state.draft_answers.get("objectifs"):
            _obj = (st.session_state.get("_objectifs_cache") or
                    st.session_state.answers.get("objectifs") or [])
            if _obj:
                st.session_state.draft_answers["objectifs"] = list(_obj)
        if not get_draft("objectifs"):
            return False, "⚠️ Retourne à la page Objectifs et sélectionne au moins un objectif pour continuer."

    for key in keys:
        if f"w_{key}" in st.session_state:
            st.session_state.draft_answers[key] = st.session_state[f"w_{key}"]
        st.session_state.answers[key] = st.session_state.draft_answers.get(
            key, st.session_state.answers.get(key, DEFAULT_ANSWERS.get(key))
        )

    if step == 1:
        selected = list(st.session_state.answers.get("objectifs") or [])
        weights = {}
        draft_weights = st.session_state.draft_answers.get("objective_weights", {}) or {}
        for objective in selected:
            widget_key = f"w_objective_weight_{safe_key(objective)}"
            weights[objective] = st.session_state.get(
                widget_key, draft_weights.get(objective, "Important")
            )
        st.session_state.answers["objective_weights"] = weights
        st.session_state.draft_answers["objective_weights"] = weights
        st.session_state["w_objective_weights"] = weights

    a = st.session_state.answers
    objectifs = set(a.get("objectifs") or [])

    if "Transmettre l'entreprise" not in objectifs or int(a.get("nb_enfants") or 0) == 0:
        for key in ["heritier_repreneur", "autres_heritiers_actifs", "soulte_envisagee",
                    "capacite_financement_soulte", "successeur_prepare", "calendrier_transmission"]:
            set_answer_and_draft(key, UNKNOWN)

    if a.get("heritier_repreneur") != YES:
        for key in ["autres_heritiers_actifs", "volonte_non_repreneurs",
                    "soulte_envisagee", "capacite_financement_soulte"]:
            set_answer_and_draft(key, UNKNOWN)

    if a.get("conjoint_present") != YES:
        set_answer_and_draft("accord_conjoint", UNKNOWN)

    if a.get("soulte_envisagee") != YES:
        set_answer_and_draft("capacite_financement_soulte", UNKNOWN)

    if "Protéger le conjoint et les proches" not in objectifs or a.get("conjoint_present") != YES:
        for key in ["conjoint_dependant", "protection_conjoint_prevue", "regime_matrimonial_adapte"]:
            set_answer_and_draft(key, UNKNOWN)

    if a.get("pacte_dutreil") != YES:
        for key in ["holding_animatrice", "audit_dutreil", "suivi_engagements"]:
            set_answer_and_draft(key, UNKNOWN)

    st.session_state.validated_steps.add(step)
    st.session_state.last_validation[step] = datetime.now().strftime("%H:%M:%S")
    return True, f"Étape {step} validée. Les réponses sont enregistrées dans le diagnostic."


def yes_no_unknown(label: str, key: str, options: List[str] | None = None, horizontal: bool = True) -> Any:
    sync_widget_from_draft(key)
    opts = options or [UNKNOWN, YES, NO]
    return st.radio(label, opts, key=f"w_{key}", horizontal=horizontal,
                    on_change=sync_draft_key, args=(key,))


def number_input(label: str, key: str, **kwargs: Any) -> Any:
    sync_widget_from_draft(key)
    return st.number_input(label, key=f"w_{key}", on_change=sync_draft_key, args=(key,), **kwargs)


def slider(label: str, key: str, min_value: int, max_value: int) -> Any:
    sync_widget_from_draft(key)
    return st.slider(label, min_value, max_value, key=f"w_{key}",
                     on_change=sync_draft_key, args=(key,))


def objective_weight_buttons(objective: str) -> None:
    """Affiche la pondération via radio — PAS de on_change pour éviter tout conflit nav."""
    widget_key = f"w_objective_weight_{safe_key(objective)}"
    saved = (
        (st.session_state.draft_answers.get("objective_weights") or {}).get(objective)
        or (st.session_state.answers.get("objective_weights") or {}).get(objective)
        or "Important"
    )
    if widget_key not in st.session_state:
        st.session_state[widget_key] = saved

    st.markdown(f"**{objective}**")
    idx = OBJECTIVE_WEIGHT_OPTIONS.index(st.session_state[widget_key]) \
        if st.session_state[widget_key] in OBJECTIVE_WEIGHT_OPTIONS else 0
    st.radio(
        "Importance",
        OBJECTIVE_WEIGHT_OPTIONS,
        index=idx,
        key=widget_key,
        horizontal=True,
        label_visibility="collapsed",
        # Pas de on_change : valeur lue par _flush_widgets_to_draft au moment de naviguer
    )


def selectbox(label: str, key: str, options: List[str]) -> Any:
    sync_widget_from_draft(key)
    current = st.session_state.get(
        f"w_{key}", st.session_state.draft_answers.get(key, DEFAULT_ANSWERS.get(key))
    )
    if current not in options:
        st.session_state[f"w_{key}"] = options[0] if options else None
        st.session_state.draft_answers[key] = st.session_state[f"w_{key}"]
    return st.selectbox(label, options, key=f"w_{key}", on_change=sync_draft_key, args=(key,))


def text_input_field(label: str, key: str, placeholder: str = "") -> Any:
    sync_widget_from_draft(key)
    return st.text_input(label, key=f"w_{key}", placeholder=placeholder,
                         on_change=sync_draft_key, args=(key,))


def text_area_field(label: str, key: str, placeholder: str = "", height: int = 90) -> Any:
    sync_widget_from_draft(key)
    return st.text_area(label, key=f"w_{key}", placeholder=placeholder, height=height,
                        on_change=sync_draft_key, args=(key,))


# =============================================================================
# 3. Scoring
# =============================================================================

def is_yes(a: Dict, key: str) -> bool:
    return a.get(key) == YES


def is_no(a: Dict, key: str) -> bool:
    return a.get(key) == NO


def is_known(a: Dict, key: str) -> bool:
    return a.get(key) not in [UNKNOWN, NA, None, ""]


def add_points(points: Dict[str, int], evidence: Dict[str, List[str]], risk: str,
               value: int, reason: str) -> None:
    if value <= 0:
        return
    points[risk] += value
    evidence.setdefault(risk, []).append(f"+{value} — {reason}")


def clamp(value: int, low: int = 0, high: int = 5) -> int:
    return max(low, min(high, value))


def niveau(score: int) -> str:
    if score == 0:
        return "Inexistant"
    if score >= 16:
        return "Critique"
    if score >= 10:
        return "Élevé"
    if score >= 5:
        return "Moyen"
    return "Faible"


def score_to_label(value: int) -> str:
    labels = {0: "Aucun signal", 1: "Faible", 2: "Moyen", 3: "Marqué", 4: "Fort", 5: "Très fort"}
    return labels.get(value, str(value))


def calculate_risks(a: Dict) -> Tuple[pd.DataFrame, Dict[str, List[str]]]:
    points = {code: 0 for code in RISK_DEFINITIONS}
    evidence: Dict[str, List[str]] = {code: [] for code in RISK_DEFINITIONS}

    objectifs = set(a.get("objectifs") or [])
    nb_enfants = int(a.get("nb_enfants") or 0)
    poids = int(a.get("poids_entreprise") or 0)
    valeur = int(a.get("valeur_entreprise") or 0)

    # ── Conserver le contrôle familial ───────────────────────────────────────
    if "Conserver le contrôle familial" in objectifs:
        if nb_enfants >= 2:
            add_points(points, evidence, "dilution", 3,
                       "Plusieurs héritiers potentiels augmentent le risque de dispersion du capital")
        if is_no(a, "clauses_entree_sortie"):
            add_points(points, evidence, "dilution", 2,
                       "Absence de clauses d'entrée/sortie des titres")
            add_points(points, evidence, "tiers", 2,
                       "Pas de clause de préemption ou d'agrément pour bloquer l'entrée de tiers")
        elif a.get("clauses_entree_sortie") == UNCERTAIN:
            add_points(points, evidence, "tiers", 1,
                       "Clauses d'entrée/sortie incertaines")
        if is_no(a, "gouvernance_formalisee"):
            add_points(points, evidence, "blocage_gouvernance", 2,
                       "Gouvernance non formalisée : risque de paralysie décisionnelle")
        elif a.get("gouvernance_formalisee") == UNCERTAIN:
            add_points(points, evidence, "blocage_gouvernance", 1,
                       "Gouvernance de la holding incertaine")
        if is_yes(a, "associes_actifs_passifs"):
            add_points(points, evidence, "blocage_gouvernance", 2,
                       "Présence d'associés actifs et passifs sans règles claires")
        if is_no(a, "politique_dividendes_definie"):
            add_points(points, evidence, "blocage_gouvernance", 1,
                       "Absence de politique de distribution définie")
        if a.get("famille_recomposee") == YES:
            add_points(points, evidence, "tiers", 1,
                       "Famille recomposée : complexité accrue des droits patrimoniaux")

    # ── Transmettre l'entreprise ──────────────────────────────────────────────
    if "Transmettre l'entreprise" in objectifs:
        if nb_enfants == 0:
            add_points(points, evidence, "successeur", 3,
                       "Aucun enfant renseigné alors que la transmission est un objectif : "
                       "successeur non identifié")
        else:
            if a.get("heritier_repreneur") == NO:
                add_points(points, evidence, "successeur", 3,
                           "Aucun héritier repreneur identifié")
            elif a.get("heritier_repreneur") == UNCERTAIN:
                add_points(points, evidence, "successeur", 2,
                           "Héritier repreneur incertain")
            if a.get("heritier_repreneur") == YES:
                if is_no(a, "successeur_prepare"):
                    add_points(points, evidence, "successeur", 2,
                               "Le successeur n'est pas encore préparé à reprendre la direction")
                elif a.get("successeur_prepare") == UNCERTAIN:
                    add_points(points, evidence, "successeur", 1,
                               "Préparation du successeur incertaine")
                if is_no(a, "calendrier_transmission"):
                    add_points(points, evidence, "successeur", 1,
                               "Pas de calendrier de transmission du pouvoir")
        if is_yes(a, "entreprise_dependante_dirigeant"):
            add_points(points, evidence, "successeur", 2,
                       "L'entreprise est fortement dépendante du dirigeant actuel")

        if nb_enfants >= 2 and a.get("heritier_repreneur") == YES:
            if is_no(a, "autres_heritiers_actifs"):
                add_points(points, evidence, "conflit_heritiers", 3,
                           "Les autres héritiers ne sont pas impliqués dans l'entreprise")
            if a.get("volonte_non_repreneurs") in ["Sortir du capital", "Incertain / non abordé"]:
                add_points(points, evidence, "conflit_heritiers", 2,
                           f"Volonté des non repreneurs : {a.get('volonte_non_repreneurs', '?')}")
            if a.get("soulte_envisagee") == YES and is_no(a, "capacite_financement_soulte"):
                add_points(points, evidence, "liquidite", 3,
                           "Soulte envisagée mais capacité de financement non validée")
            elif a.get("soulte_envisagee") == YES:
                add_points(points, evidence, "liquidite", 1,
                           "Soulte envisagée : vérifier la faisabilité du financement")

        if nb_enfants >= 2 and is_no(a, "dialogue_familial"):
            add_points(points, evidence, "conflit_heritiers", 2,
                       "Aucun dialogue familial n'a encore été organisé")

    # ── Optimiser la fiscalité ────────────────────────────────────────────────
    if "Optimiser la fiscalité" in objectifs:
        if is_no(a, "simulation_fiscale"):
            add_points(points, evidence, "fiscalite", 3,
                       "Aucune simulation fiscale réalisée ou prévue")
        elif a.get("simulation_fiscale") == UNCERTAIN:
            add_points(points, evidence, "fiscalite", 1,
                       "Simulation fiscale prévue mais non encore réalisée")
        if a.get("pacte_dutreil") == NO:
            add_points(points, evidence, "dutreil", 1,
                       "Pacte Dutreil non envisagé malgré l'objectif fiscal")
        elif a.get("pacte_dutreil") == YES:
            if a.get("holding_animatrice") in [NO, UNCERTAIN]:
                add_points(points, evidence, "dutreil", 3,
                           "Qualification de holding animatrice non établie ou incertaine")
            if is_no(a, "audit_dutreil"):
                add_points(points, evidence, "dutreil", 2,
                           "Aucun audit Dutreil réalisé")
            if is_no(a, "suivi_engagements"):
                add_points(points, evidence, "dutreil", 2,
                           "Pas de suivi des engagements de conservation")

    # ── Protéger le conjoint et les proches ──────────────────────────────────
    if "Protéger le conjoint et les proches" in objectifs or a.get("conjoint_present") == YES:
        if a.get("conjoint_present") == YES:
            if is_yes(a, "conjoint_dependant") and is_no(a, "protection_conjoint_prevue"):
                add_points(points, evidence, "conjoint", 4,
                           "Conjoint financièrement dépendant sans dispositif de protection prévu")
            elif is_yes(a, "conjoint_dependant"):
                add_points(points, evidence, "conjoint", 2,
                           "Conjoint financièrement dépendant du dirigeant")
            elif is_no(a, "protection_conjoint_prevue"):
                add_points(points, evidence, "conjoint", 2,
                           "Aucun dispositif de protection du conjoint prévu")
            if is_no(a, "regime_matrimonial_adapte"):
                add_points(points, evidence, "conjoint", 1,
                           "Régime matrimonial non analysé ou non adapté")
            if a.get("famille_recomposee") == YES:
                add_points(points, evidence, "conjoint", 1,
                           "Famille recomposée : droits du conjoint à sécuriser")
            if a.get("accord_conjoint") == NO:
                add_points(points, evidence, "conjoint", 2,
                           "Le conjoint n'adhère pas au projet de transmission")
            elif a.get("accord_conjoint") == UNCERTAIN:
                add_points(points, evidence, "conjoint", 1,
                           "L'adhésion du conjoint au projet est incertaine")

    # ── Préserver l'équité entre les héritiers ───────────────────────────────
    if "Préserver l'équité entre les héritiers" in objectifs:
        if nb_enfants >= 2:
            if is_no(a, "valorisation_independante"):
                add_points(points, evidence, "contestation", 2,
                           "Aucune valorisation indépendante des titres prévue")
            if is_no(a, "audit_civil"):
                add_points(points, evidence, "contestation", 2,
                           "Aucun audit civil et successoral réalisé avec le notaire")
        if a.get("famille_recomposee") == YES:
            add_points(points, evidence, "contestation", 2,
                       "Famille recomposée : risque de contestation successorale renforcé")
            add_points(points, evidence, "conflit_heritiers", 2,
                       "Famille recomposée : risque de divergence d'intérêts entre héritiers renforcé")

    # ── Diversifier le patrimoine ─────────────────────────────────────────────
    if "Diversifier le patrimoine" in objectifs or poids >= 40:
        if poids >= 70:
            add_points(points, evidence, "dependance", 3,
                       f"Poids de l'entreprise très élevé ({poids}%) : risque de concentration critique")
        elif poids >= 40:
            add_points(points, evidence, "dependance", 2,
                       f"Poids de l'entreprise élevé ({poids}%) : diversification insuffisante")
        if is_no(a, "diversification"):
            add_points(points, evidence, "dependance", 2,
                       "Aucune diversification patrimoniale organisée")
        if a.get("actifs_liquides") == "Faible":
            add_points(points, evidence, "dependance", 1,
                       "Faible niveau d'actifs liquides hors entreprise")
        if a.get("endettement_familial") == "Élevé":
            add_points(points, evidence, "dependance", 1,
                       "Endettement familial élevé")
        if is_no(a, "prevoyance"):
            add_points(points, evidence, "dependance", 1,
                       "Pas de prévoyance ou assurance décès prévue")
        if a.get("besoin_revenus_famille") == "Élevé":
            add_points(points, evidence, "liquidite", 2,
                       "Besoin de revenus réguliers élevé : risque de pression sur les dividendes")
        elif a.get("besoin_revenus_famille") == "Moyen":
            add_points(points, evidence, "liquidite", 1,
                       "Besoin de revenus réguliers moyen à surveiller")

    # ── Règles transversales ──────────────────────────────────────────────────
    if a.get("urgence_evenement") == YES:
        for risk in ["successeur", "fiscalite", "conjoint"]:
            add_points(points, evidence, risk, 1,
                       "Événement d'urgence ou de fragilité signalé")

    if a.get("delai_transmission") == "Moins de 12 mois":
        add_points(points, evidence, "fiscalite", 1,
                   "Horizon de transmission inférieur à 12 mois : peu de temps pour optimiser")
        add_points(points, evidence, "successeur", 1,
                   "Transmission imminente : préparer rapidement le successeur")

    if is_no(a, "suivi_annuel"):
        add_points(points, evidence, "suivi", 2,
                   "Pas de revue annuelle de la stratégie patrimoniale prévue")
    if is_no(a, "formalisation_rapport"):
        add_points(points, evidence, "suivi", 1,
                   "Pas de rapport écrit de diagnostic prévu")

    if a.get("famille_recomposee") == YES:
        if "Préserver l'équité entre les héritiers" not in objectifs:
            add_points(points, evidence, "contestation", 1,
                       "Famille recomposée : risque successoral à ne pas négliger hors objectif déclaré")
        if "Protéger le conjoint et les proches" not in objectifs:
            add_points(points, evidence, "conjoint", 1,
                       "Famille recomposée : droits du conjoint à vérifier")

    # Valeur >= 5M€ — signaux autonomes
    if valeur >= 5_000_000:
        add_points(points, evidence, "fiscalite", 1,
                   f"Valeur élevée ({valeur/1e6:.1f}M€) : enjeu fiscal même hors objectif déclaré")
        if a.get("pacte_dutreil") != YES:
            add_points(points, evidence, "dutreil", 1,
                       f"Valeur >= 5M€ : le Pacte Dutreil mérite d'être examiné")
        add_points(points, evidence, "contestation", 1,
                   f"Valeur >= 5M€ : risque de contestation accru par l'importance des enjeux")

    return build_df(points, evidence, a)


def build_df(points: Dict[str, int], evidence: Dict[str, List[str]], answers: Dict | None = None) -> Tuple[pd.DataFrame, Dict]:
    rows = []
    for code, definition in RISK_DEFINITIONS.items():
        raw_prob = clamp(points.get(code, 0))
        weight_label, factor = weight_for_risk(code, answers)
        weighted_prob = raw_prob * factor
        grav = definition.gravite_base
        score = round(weighted_prob * grav, 1)
        niv = niveau(int(score))
        rows.append({
            "Code": code,
            "Objectif": definition.objectif,
            "Risque": definition.libelle,
            "Probabilité": raw_prob,
            "Gravité": grav,
            "Score": score,
            "Niveau": niv,
            "Pondération objectif": weight_label,
            "Outils à étudier": ", ".join(definition.outils),
            "Actions préventives": ", ".join(definition.actions_preventives),
            "Professionnels": ", ".join(definition.professionnels),
            "Signaux retenus": "\n".join(evidence.get(code, [])),
        })
    df = pd.DataFrame(rows)
    df = df.sort_values(["Score", "Probabilité", "Gravité"], ascending=False).reset_index(drop=True)
    return df, evidence


def missing_data(a: Dict) -> List[str]:
    gaps: List[str] = []
    objectifs = set(a.get("objectifs") or [])
    if not objectifs:
        gaps.append("Aucun objectif sélectionné — le scoring ne peut pas s'activer.")
        return gaps
    if not a.get("client_name"):
        gaps.append("Nom du client non renseigné.")
    if a.get("nb_enfants") in [UNKNOWN, None]:
        gaps.append("Nombre d'enfants non renseigné.")
    if "Transmettre l'entreprise" in objectifs and a.get("heritier_repreneur") == UNKNOWN:
        gaps.append("Présence d'un repreneur non renseignée (objectif transmission sélectionné).")
    if a.get("conjoint_present") == UNKNOWN:
        gaps.append("Présence d'un conjoint non renseignée.")
    if int(a.get("valeur_entreprise") or 0) == 0:
        gaps.append("Valeur de l'entreprise non renseignée (scoring fiscal et liquidité imprécis).")
    if int(a.get("poids_entreprise") or 0) == 0 and "Diversifier le patrimoine" in objectifs:
        gaps.append("Poids de l'entreprise dans le patrimoine non renseigné.")
    if "Optimiser la fiscalité" in objectifs and a.get("simulation_fiscale") == UNKNOWN:
        gaps.append("Simulation fiscale non renseignée.")
    if a.get("pacte_dutreil") == YES and a.get("holding_animatrice") == UNKNOWN:
        gaps.append("Qualification de holding animatrice non renseignée (Pacte Dutreil en cours).")
    return gaps


# =============================================================================
# 4. Helpers d'affichage (HTML, cartes)
# =============================================================================

def priority_badge(level: str) -> str:
    color = PRIORITY_COLORS.get(level, "#6b7280")
    return f'<span class="badge" style="background:{color}">{level}</span>'


def card(title: str, body: str, accent: str = "#0ea5e9") -> None:
    st.markdown(
        f'<div class="card" style="border-left:5px solid {accent}">'
        f'<h3 style="color:{accent}!important">{title}</h3>'
        f'{body}'
        f'</div>',
        unsafe_allow_html=True,
    )


def render_list(items: List[str]) -> str:
    if not items:
        return "<p><em>—</em></p>"
    li = "".join(f"<li>{item}</li>" for item in items)
    return f"<ul>{li}</ul>"


# =============================================================================
# 5. Justifications des outils
# =============================================================================

_TOOL_JUSTIFICATIONS: Dict[str, Dict[str, str]] = {
    "dilution": {
        "Holding familiale": "Regroupe les titres dans une structure commune, permettant de centraliser le contrôle et d'organiser la transmission par étapes.",
        "Clause d'agrément": "Soumet toute cession de titres à l'accord préalable des associés pour éviter la dilution involontaire.",
        "Clause de préemption": "Donne aux associés existants la priorité pour racheter les titres mis en vente.",
        "Actions de préférence": "Permettent de dissocier droits économiques et droits de vote pour préserver le contrôle.",
        "Pacte d'associés": "Organise les relations entre associés, la gouvernance et les règles de cession dans un document contractuel.",
    },
    "tiers": {
        "Clause d'agrément": "Outil central pour soumettre toute entrée d'un tiers à l'approbation des associés.",
        "Clause de préemption": "Droit de rachat prioritaire permettant de bloquer la sortie vers un tiers non souhaité.",
        "Pacte d'associés": "Précise l'ensemble des règles d'entrée et de sortie pour sécuriser le capital familial.",
        "Statuts adaptés": "Les statuts de SAS offrent une grande flexibilité pour verrouiller la gouvernance.",
    },
    "blocage_gouvernance": {
        "Pacte d'associés": "Document clé pour définir les règles de majorité, de blocage et de sortie entre associés.",
        "Charte familiale": "Engage la famille autour de valeurs communes et de règles de gouvernance non juridiques.",
        "Conseil de famille": "Instance de dialogue régulière entre membres de la famille actionnaire.",
        "Comité stratégique": "Structure associant les non-actifs aux décisions importantes sans droit de vote.",
        "Statuts de SAS": "Structure très flexible permettant d'organiser la gouvernance sur mesure.",
    },
    "successeur": {
        "Gouvernance transitoire": "Assure la continuité opérationnelle pendant la phase de transmission progressive.",
        "Direction externe temporaire": "Pallie l'absence de successeur identifié ou préparé.",
        "Actions de préférence": "Permettent de séparer le contrôle opérationnel de la propriété économique.",
        "Pacte d'associés": "Organise la transmission progressive des pouvoirs et les droits du successeur.",
        "Calendrier de transmission": "Outil de planification pour anticiper et formaliser les étapes clés.",
    },
    "conflit_heritiers": {
        "Donation-partage": "Acte notarié permettant de répartir les biens de son vivant, avec accord des parties.",
        "Soulte": "Compensation financière versée aux héritiers qui ne reprennent pas l'entreprise.",
        "Pacte d'associés": "Encadre les droits et devoirs des associés repreneurs et non repreneurs.",
        "Mécanisme de sortie": "Organise la liquidité des non repreneurs pour éviter tout blocage.",
        "Droits financiers différenciés": "Permettent de distinguer droits de vote et droits aux dividendes selon le rôle.",
    },
    "contestation": {
        "Donation-partage": "Cristallise la valeur au moment de la donation, réduisant les risques de requalification.",
        "Audit civil": "Analyse des droits successoraux pour identifier en amont les risques de contestation.",
        "Valorisation indépendante": "Expertise indépendante de la valeur des titres pour éviter les litiges sur l'évaluation.",
        "Démembrement": "Technique permettant de transmettre la nue-propriété en conservant l'usufruit.",
        "Testament": "Acte formel exprimant la volonté du testateur et réduisant les ambiguïtés successorales.",
    },
    "liquidite": {
        "Soulte échelonnée": "Paiement de la compensation en plusieurs fois pour alléger la trésorerie du repreneur.",
        "Family Buy-Out": "Rachat de l'entreprise par la famille, structuré avec un effet de levier financier.",
        "Paiement différé/fractionné": "Option légale permettant d'étaler les droits de mutation dans le temps.",
        "Assurance-vie": "Instrument de liquidité disponible pour financer soultes ou droits de succession.",
        "Politique de distribution encadrée": "Définit les règles de dividendes pour assurer la liquidité des associés non actifs.",
    },
    "fiscalite": {
        "Pacte Dutreil": "Dispositif central permettant une réduction de 75% de la base taxable lors de la transmission.",
        "Donation anticipée": "Transmettre de son vivant, profitant des abattements et d'une valorisation possiblement plus basse.",
        "Démembrement": "Sépare usufruit et nue-propriété, réduisant la valeur imposable transmise.",
        "Paiement différé et fractionné": "Permet d'étaler le paiement des droits de mutation sur plusieurs années.",
    },
    "dutreil": {
        "Audit Dutreil": "Vérifie la conformité aux conditions légales du Pacte Dutreil avant et pendant sa durée.",
        "Documentation holding animatrice": "Établit et documente la qualification de holding animatrice, condition essentielle.",
        "Suivi des engagements": "Contrôle annuel du respect des engagements de conservation des titres.",
        "Revue juridique annuelle": "Assure le maintien des conditions d'éligibilité au fil du temps et des évolutions législatives.",
    },
    "conjoint": {
        "Assurance-vie": "Outil de protection du conjoint survivant et d'optimisation successorale.",
        "Régime matrimonial adapté": "Ajuster le contrat de mariage pour mieux protéger les droits patrimoniaux du conjoint.",
        "Donation au dernier vivant": "Avantage le conjoint survivant au-delà des droits légaux minimaux.",
        "SCI familiale": "Permet de transmettre le patrimoine immobilier en protégeant le logement familial.",
        "Mandat de protection future": "Anticipe une éventuelle incapacité du dirigeant et protège le conjoint.",
    },
    "dependance": {
        "Immobilier locatif": "Diversification patrimoniale classique générant des revenus réguliers.",
        "Contrats d'assurance-vie": "Outil de diversification liquide, fiscalement avantageux, hors succession.",
        "SCPI/OPCI": "Investissement immobilier mutualisé sans contrainte de gestion directe.",
        "Prévoyance dirigeant": "Assurance couvrant les risques décès, invalidité et incapacité du dirigeant.",
        "Plan d'épargne retraite": "Diversification vers des actifs financiers avec avantage fiscal à l'entrée.",
    },
    "suivi": {
        "Tableau de bord patrimonial": "Suivi annuel structuré de l'évolution du patrimoine et des risques identifiés.",
        "Revue juridique annuelle": "Vérification annuelle de la conformité des dispositifs mis en place.",
        "Rapport de diagnostic CGP": "Document de synthèse annuel remis au client pour formaliser les recommandations.",
        "Réunion annuelle associés": "Instance de gouvernance régulière pour maintenir le dialogue et la cohésion familiale.",
    },
}


def get_tool_justifications(code: str) -> Dict[str, str]:
    return _TOOL_JUSTIFICATIONS.get(code, {})


def render_tool_justifications(code: str) -> str:
    justifications = get_tool_justifications(code)
    if not justifications:
        return "<p><em>Justifications non disponibles.</em></p>"
    rows = "".join(
        f"<tr><td><strong>{tool}</strong></td><td>{just}</td></tr>"
        for tool, just in justifications.items()
    )
    return (
        f'<table class="justification-table"><thead>'
        f"<tr><th>Outil</th><th>Pourquoi cet outil</th></tr>"
        f"</thead><tbody>{rows}</tbody></table>"
    )


# =============================================================================
# 6. Plan d'action
# =============================================================================

def build_action_plan(df: pd.DataFrame) -> pd.DataFrame:
    """Génère un plan d'action priorisé depuis le DataFrame des risques."""
    if df.empty:
        return pd.DataFrame(
            columns=["Horizon", "Priorité", "Action", "Risques concernés", "Professionnels"]
        )

    rows = []

    def _add_group(horizon: str, niveaux: List[str]) -> None:
        seen: Dict[str, Dict] = {}
        sub = df[df["Niveau"].isin(niveaux)].copy()
        for _, row in sub.iterrows():
            definition = RISK_DEFINITIONS[row["Code"]]
            for action in definition.actions_preventives[:2]:
                key = action[:60]
                if key not in seen:
                    seen[key] = {
                        "Horizon": horizon,
                        "Priorité": row["Niveau"],
                        "Action": action,
                        "Risques": [definition.libelle],
                        "Pros": list(definition.professionnels),
                    }
                else:
                    if definition.libelle not in seen[key]["Risques"]:
                        seen[key]["Risques"].append(definition.libelle)
                    for p in definition.professionnels:
                        if p not in seen[key]["Pros"]:
                            seen[key]["Pros"].append(p)
        for ag in seen.values():
            rows.append({
                "Horizon": ag["Horizon"],
                "Priorité": ag["Priorité"],
                "Action": ag["Action"],
                "Risques concernés": ", ".join(ag["Risques"]),
                "Professionnels": ", ".join(ag["Pros"]),
            })

    _add_group("Actions immédiates", ["Critique", "Élevé"])
    _add_group("Actions à moyen terme", ["Moyen"])

    rows += [
        {
            "Horizon": "Suivi annuel",
            "Priorité": "Faible",
            "Action": "Mettre à jour le diagnostic avec le client chaque année",
            "Risques concernés": "Tous les risques identifiés",
            "Professionnels": "CGP",
        },
        {
            "Horizon": "Suivi annuel",
            "Priorité": "Faible",
            "Action": "Vérifier la conformité des engagements Dutreil si applicable",
            "Risques concernés": "Remise en cause du Pacte Dutreil",
            "Professionnels": "Avocat fiscaliste, Expert-comptable",
        },
    ]

    if not rows:
        return pd.DataFrame(
            columns=["Horizon", "Priorité", "Action", "Risques concernés", "Professionnels"]
        )

    plan_df = pd.DataFrame(rows)
    h_order = {"Actions immédiates": 0, "Actions à moyen terme": 1, "Suivi annuel": 2}
    p_order = {"Critique": 0, "Élevé": 1, "Moyen": 2, "Faible": 3, "Inexistant": 4}
    plan_df["_h"] = plan_df["Horizon"].map(h_order).fillna(3)
    plan_df["_p"] = plan_df["Priorité"].map(p_order).fillna(5)
    plan_df = (
        plan_df.sort_values(["_h", "_p"])
        .drop(columns=["_h", "_p"])
        .reset_index(drop=True)
    )
    return plan_df


# =============================================================================
# 7. Export Markdown
# =============================================================================

def markdown_report(client_name: str, answers: Dict, df: pd.DataFrame) -> str:
    date_str = datetime.now().strftime("%d/%m/%Y")
    detected = df[df["Score"] > 0].copy()

    out = [
        f"# Diagnostic Holding Familiale — {client_name or 'Client'}",
        f"*Généré le {date_str} — Système expert CGP*",
        "",
        "---",
        "",
        "## Contexte du dossier",
        f"- **Entreprise** : {answers.get('company_name', 'N/A')}",
        f"- **Forme juridique** : {answers.get('company_form', 'N/A')}",
        f"- **Activité** : {answers.get('company_activity', 'N/A')}",
        f"- **Âge du dirigeant** : {answers.get('client_age', 'N/A')}",
        f"- **CGP / cabinet** : {answers.get('cgp_name', 'N/A')}",
        f"- **Maturité du projet** : {answers.get('maturite_projet', 'N/A')}",
        f"- **Horizon de transmission** : {answers.get('delai_transmission', 'N/A')}",
        "",
        "## Objectifs du dirigeant",
    ]
    objectifs = answers.get("objectifs") or []
    for obj in objectifs:
        weight = (answers.get("objective_weights") or {}).get(obj, "Important")
        out.append(f"- {obj} *(pondération : {weight})*")
    if answers.get("objectif_libre"):
        out += ["", f"**Objectif en propres termes** : *{answers['objectif_libre']}*"]
    out += [
        "",
        "---",
        "",
        f"## Résultats — {len(detected)} risque(s) identifié(s)",
        "",
        "| Risque | Score | Niveau |",
        "|--------|-------|--------|",
    ]
    for _, row in detected.iterrows():
        out.append(f"| {row['Risque']} | {int(row['Score'])} | **{row['Niveau']}** |")
    out += ["", "---", "", "## Analyse détaillée par risque", ""]
    for _, row in detected.iterrows():
        definition = RISK_DEFINITIONS[row["Code"]]
        out += [
            f"### {row['Risque']} — {row['Niveau']} (Score : {int(row['Score'])})",
            f"*Objectif concerné : {definition.objectif}*",
            "",
            "**Outils à étudier :** " + ", ".join(definition.outils),
            "",
            "**Actions préventives :**",
        ]
        for action in definition.actions_preventives:
            out.append(f"- {action}")
        out += [
            "",
            "**Professionnels à mobiliser :** " + ", ".join(definition.professionnels),
            "",
        ]
    out += ["---", "", "## Plan d'action", ""]
    plan = build_action_plan(df)
    if not plan.empty:
        for horizon in ["Actions immédiates", "Actions à moyen terme", "Suivi annuel"]:
            hdf = plan[plan["Horizon"] == horizon]
            if not hdf.empty:
                out.append(f"### {horizon}")
                for _, row in hdf.iterrows():
                    out.append(
                        f"- [{row['Priorité']}] {row['Action']} "
                        f"*(Pros : {row['Professionnels']})*"
                    )
                out.append("")
    if answers.get("observations"):
        out += ["---", "", "## Observations du CGP", "", answers["observations"], ""]
    out += ["---", "", "*Document généré automatiquement — Système expert CGP Holding Familiale*"]
    return "\n".join(out)


# =============================================================================
# 8. Export Word (DOCX)
# =============================================================================

def create_docx_report(client_name: str, answers: Dict, df: pd.DataFrame, evidence: Dict) -> bytes:
    if Document is None:
        raise RuntimeError("La dépendance python-docx n'est pas installée.")

    detected = df[df["Score"] > 0].copy()
    doc = Document()

    # ── Helpers internes ───────────────────────────────────────────────────
    def shade_cell(cell, hex_color: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color.lstrip("#"))
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 11,
                       color_hex: str = "000000") -> None:
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(text))
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(color_hex)

    def add_bullets(items: List[str]) -> None:
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item).font.size = Pt(10)

    def add_small_note(text: str) -> None:
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(9)
        p.runs[0].italic = True

    def add_heading(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    # ── Couverture ────────────────────────────────────────────────────────
    doc.add_heading("Diagnostic Holding Familiale", level=0)
    doc.add_paragraph(f"Client : {client_name or 'N/A'}")
    doc.add_paragraph(f"Entreprise : {answers.get('company_name', 'N/A')}")
    doc.add_paragraph(f"CGP / Cabinet : {answers.get('cgp_name', 'N/A')}")
    doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph(
        "Document confidentiel — généré automatiquement par le système expert CGP Holding Familiale"
    )
    doc.add_page_break()

    # ── Synthèse exécutive ────────────────────────────────────────────────
    add_heading("Synthèse exécutive", 1)
    n_crit = int((detected["Niveau"] == "Critique").sum()) if not detected.empty else 0
    n_elev = int((detected["Niveau"] == "Élevé").sum()) if not detected.empty else 0
    doc.add_paragraph(
        f"Ce diagnostic identifie {len(detected)} risque(s) actif(s) sur 12 analysés, "
        f"dont {n_crit} critique(s) et {n_elev} élevé(s)."
    )
    orientation = answers.get("rapport_orientation", "Équilibré")
    doc.add_paragraph(f"Orientation du rapport : {orientation}.")

    # ── Contexte ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading("Contexte du dossier", 1)
    fields = [
        ("Âge du dirigeant", answers.get("client_age", "N/A")),
        ("Entreprise", answers.get("company_name", "N/A")),
        ("Forme juridique", answers.get("company_form", "N/A")),
        ("Activité / secteur", answers.get("company_activity", "N/A")),
        ("Maturité du projet", answers.get("maturite_projet", "N/A")),
        ("Horizon de transmission", answers.get("delai_transmission", "N/A")),
        ("Valeur de l'entreprise", f"{int(answers.get('valeur_entreprise', 0)):,} €".replace(",", " ")),
        ("Poids dans le patrimoine", f"{answers.get('poids_entreprise', 0)} %"),
        ("Nombre d'enfants", answers.get("nb_enfants", "N/A")),
        ("Conjoint présent", answers.get("conjoint_present", "N/A")),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)

    # ── Objectifs et pondérations ─────────────────────────────────────────
    doc.add_page_break()
    add_heading("Objectifs et pondérations", 1)
    objectifs = answers.get("objectifs") or []
    if objectifs:
        weights = answers.get("objective_weights") or {}
        table_obj = doc.add_table(rows=1 + len(objectifs), cols=2)
        table_obj.style = "Table Grid"
        table_obj.rows[0].cells[0].text = "Objectif"
        table_obj.rows[0].cells[1].text = "Pondération"
        for idx, obj in enumerate(objectifs, start=1):
            table_obj.rows[idx].cells[0].text = obj
            table_obj.rows[idx].cells[1].text = weights.get(obj, "Important")
    if answers.get("objectif_libre"):
        doc.add_paragraph(f"\nObjectif exprimé par le client : {answers['objectif_libre']}")

    # ── Matrice des risques ───────────────────────────────────────────────
    doc.add_page_break()
    add_heading("Matrice des risques", 1)
    cols_matrix = ["Risque", "Niveau", "Score", "Probabilité", "Gravité"]
    t = doc.add_table(rows=1 + len(df), cols=len(cols_matrix))
    t.style = "Table Grid"
    LEVEL_COLORS = {"Critique": "7f1d1d", "Élevé": "b45309",
                    "Moyen": "1d4ed8", "Faible": "047857", "Inexistant": "6b7280"}
    for j, col in enumerate(cols_matrix):
        t.rows[0].cells[j].text = col
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        t.rows[i].cells[0].text = str(row["Risque"])
        set_cell_text(t.rows[i].cells[1], str(row["Niveau"]), bold=True,
                      color_hex=LEVEL_COLORS.get(str(row["Niveau"]), "000000"))
        t.rows[i].cells[2].text = str(int(row["Score"]))
        t.rows[i].cells[3].text = str(int(row["Probabilité"]))
        t.rows[i].cells[4].text = str(int(row["Gravité"]))

    # ── Analyse détaillée ─────────────────────────────────────────────────
    if not detected.empty:
        doc.add_page_break()
        add_heading("Analyse détaillée des risques identifiés", 1)
        for _, row in detected.iterrows():
            definition = RISK_DEFINITIONS[row["Code"]]
            add_heading(f"{row['Risque']} — {row['Niveau']} (Score : {int(row['Score'])})", 2)
            doc.add_paragraph(f"Objectif concerné : {definition.objectif}")
            doc.add_paragraph(
                f"Score : {int(row['Score'])} | "
                f"Probabilité : {score_to_label(int(row['Probabilité']))} | "
                f"Gravité : {score_to_label(int(row['Gravité']))}"
            )
            signals = evidence.get(row["Code"], [])
            if signals:
                add_heading("Signaux retenus", 3)
                add_bullets(signals)
            add_heading("Conséquences possibles", 3)
            add_bullets(definition.consequences)
            add_heading("Outils à étudier", 3)
            add_bullets(definition.outils)
            add_heading("Actions préventives", 3)
            add_bullets(definition.actions_preventives)
            doc.add_paragraph(
                f"Professionnels à mobiliser : {', '.join(definition.professionnels)}"
            )

    # ── Plan d'action ─────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading("Plan d'action hiérarchisé", 1)
    plan = build_action_plan(df)
    if not plan.empty:
        for horizon in ["Actions immédiates", "Actions à moyen terme", "Suivi annuel"]:
            hdf = plan[plan["Horizon"] == horizon]
            if not hdf.empty:
                add_heading(horizon, 2)
                for _, row in hdf.iterrows():
                    doc.add_paragraph(
                        f"[{row['Priorité']}] {row['Action']} "
                        f"— Pros : {row['Professionnels']}",
                        style="List Bullet",
                    )

    # ── Observations ──────────────────────────────────────────────────────
    if answers.get("observations"):
        doc.add_page_break()
        add_heading("Observations du CGP", 1)
        doc.add_paragraph(answers["observations"])

    # ── Footer ────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_paragraph(
        "Document généré automatiquement — Système expert CGP Holding Familiale\n"
        "Ce rapport est un outil d'aide à la décision. Il ne constitue pas un conseil juridique "
        "ou fiscal et doit être complété par les professionnels compétents."
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# =============================================================================
# 9. Export Excel (XLSX)
# =============================================================================

def create_xlsx_matrix(df: pd.DataFrame) -> bytes:
    """Génère une matrice Excel lisible et directement exploitable."""
    if Workbook is None:
        raise RuntimeError("La dépendance openpyxl n'est pas installée.")

    export_columns = [
        "Objectif", "Risque", "Probabilité", "Gravité", "Score", "Niveau", "Pondération objectif",
        "Signaux retenus", "Outils à étudier", "Actions préventives", "Professionnels",
    ]
    wb = Workbook()
    ws = wb.active
    ws.title = "Matrice des risques"

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(export_columns))
    title_cell = ws.cell(row=1, column=1)
    title_cell.value = "Matrice des risques – Diagnostic holding familiale"
    title_cell.font = Font(bold=True, size=16, color="FFFFFF")
    title_cell.fill = PatternFill("solid", fgColor="0B7285")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(export_columns))
    subtitle_cell = ws.cell(row=2, column=1)
    subtitle_cell.value = "Document de travail exporté depuis le prototype de système expert CGP"
    subtitle_cell.font = Font(italic=True, color="475569")
    subtitle_cell.alignment = Alignment(horizontal="center", vertical="center")

    header_row = 4
    for col_idx, col_name in enumerate(export_columns, start=1):
        cell = ws.cell(row=header_row, column=col_idx)
        cell.value = col_name
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="0B7285")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    priority_fills = {
        "Critique": "FEE2E2",
        "Élevé": "FFEDD5",
        "Moyen": "DBEAFE",
        "Faible": "DCFCE7",
        "Inexistant": "F1F5F9",
    }

    available_cols = [c for c in export_columns if c in df.columns]
    sorted_df = df[available_cols].copy()
    for row_idx, (_, row) in enumerate(sorted_df.iterrows(), start=header_row + 1):
        priority = row.get("Niveau", "")
        fill_color = priority_fills.get(priority, "FFFFFF")
        for col_idx, col_name in enumerate(available_cols, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = row.get(col_name, "")
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
            if col_name == "Niveau":
                cell.fill = PatternFill("solid", fgColor=fill_color)
                cell.font = Font(bold=True)
            elif row_idx % 2 == 0:
                cell.fill = PatternFill("solid", fgColor="F8FAFC")

    col_widths = {"A": 24, "B": 34, "C": 12, "D": 10, "E": 10, "F": 18, "G": 22,
                  "H": 55, "I": 42, "J": 48, "K": 34}
    for col_letter, width in col_widths.items():
        ws.column_dimensions[col_letter].width = width
    for row in range(header_row + 1, header_row + 1 + len(sorted_df)):
        ws.row_dimensions[row].height = 60
    ws.freeze_panes = "A5"
    ws.auto_filter.ref = (
        f"A{header_row}:{get_column_letter(len(available_cols))}{header_row + len(sorted_df)}"
    )

    # Plan d'action
    plan_df = build_action_plan(df)
    ws_plan = wb.create_sheet("Plan d'action")
    plan_columns = ["Horizon", "Priorité", "Action", "Risques concernés", "Professionnels"]
    ws_plan.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(plan_columns))
    plan_title = ws_plan.cell(row=1, column=1)
    plan_title.value = "Plan d'action hiérarchisé"
    plan_title.font = Font(bold=True, size=15, color="FFFFFF")
    plan_title.fill = PatternFill("solid", fgColor="0B7285")
    plan_title.alignment = Alignment(horizontal="center", vertical="center")
    ws_plan.row_dimensions[1].height = 28

    for col_idx, col_name in enumerate(plan_columns, start=1):
        cell = ws_plan.cell(row=3, column=col_idx)
        cell.value = col_name
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="0B7285")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    horizon_fills = {
        "Actions immédiates": "FEE2E2",
        "Actions à moyen terme": "FFEDD5",
        "Suivi annuel": "DBEAFE",
    }
    for row_idx, (_, row) in enumerate(plan_df.iterrows(), start=4):
        h_fill = horizon_fills.get(row.get("Horizon", ""), "F8FAFC")
        for col_idx, col_name in enumerate(plan_columns, start=1):
            cell = ws_plan.cell(row=row_idx, column=col_idx)
            cell.value = row.get(col_name, "")
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
            if col_name == "Horizon":
                cell.fill = PatternFill("solid", fgColor=h_fill)
                cell.font = Font(bold=True)
            elif row_idx % 2 == 0:
                cell.fill = PatternFill("solid", fgColor="F8FAFC")

    for col_letter, width in {"A": 24, "B": 16, "C": 70, "D": 36, "E": 42}.items():
        ws_plan.column_dimensions[col_letter].width = width
    for row in range(4, 4 + len(plan_df)):
        ws_plan.row_dimensions[row].height = 55
    ws_plan.freeze_panes = "A4"
    if len(plan_df) > 0:
        ws_plan.auto_filter.ref = (
            f"A3:{get_column_letter(len(plan_columns))}{3 + len(plan_df)}"
        )

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()



VIEWS = ["Questionnaire adaptatif", "Résultats et solutions", "Exporter", "Règles de décision", "Debug validation"]

# ─── Initialisation de la session (DOIT être avant tout accès à session_state) ──
init_app()

# ─── Initialisation état UI ──────────────────────────────────────────────────

st.session_state.setdefault("app_page", VIEWS[0])
st.session_state.setdefault("current_subpage", "dossier")
st.session_state.setdefault("nav_error", None)
st.session_state.setdefault("_objectifs_cache", [])

# Restaure le cache objectifs depuis draft/answers à chaque rerun
if not st.session_state.get("_objectifs_cache"):
    _obj = (st.session_state.draft_answers.get("objectifs") or
            st.session_state.answers.get("objectifs") or [])
    if _obj:
        st.session_state["_objectifs_cache"] = _obj

# ─── Sous-pages du questionnaire ─────────────────────────────────────────────

ALL_SUBPAGES: List[Dict[str, Any]] = [
    # Section 1 – Cadrage & objectifs
    {"id": "dossier",        "title": "Identification du dossier",   "icon": "📋", "step": 1,
     "desc": "Informations générales sur le dossier et le client"},
    {"id": "maturite",       "title": "Maturité & urgence",          "icon": "🗓️", "step": 1,
     "desc": "Horizon de transmission et signaux d'urgence"},
    {"id": "objectifs",      "title": "Objectifs du dirigeant",      "icon": "🎯", "step": 1,
     "desc": "Objectifs poursuivis — socle du scoring des risques"},
    {"id": "poids",          "title": "Importance des objectifs",    "icon": "⚖️", "step": 1,
     "desc": "Pondération pour affiner l'analyse des risques"},
    {"id": "rapport",        "title": "Personnalisation du rapport", "icon": "📄", "step": 1,
     "desc": "Adapter le rapport exporté au besoin du client"},
    # Section 2 – Famille
    {"id": "famille",        "title": "Situation familiale",         "icon": "👨‍👩‍👧", "step": 2,
     "desc": "Composition de la famille, enfants, conjoint"},
    {"id": "repreneur",      "title": "Repreneur & transmission",    "icon": "🔑", "step": 2,
     "desc": "Successeur identifié et organisation de la reprise"},
    {"id": "dialogue",       "title": "Dialogue familial",           "icon": "💬", "step": 2,
     "desc": "Organisation du dialogue entre héritiers"},
    {"id": "conjoint",       "title": "Protection du conjoint",      "icon": "🛡️", "step": 2,
     "desc": "Droits et protection patrimoniale du conjoint"},
    {"id": "succession_civ", "title": "Sécurisation successorale",   "icon": "⚖️", "step": 2,
     "desc": "Valorisation et audit civil pour éviter la contestation"},
    # Section 3 – Patrimoine
    {"id": "patrimoine",     "title": "Situation patrimoniale",      "icon": "💼", "step": 3,
     "desc": "Valeur, composition et liquidité du patrimoine"},
    {"id": "securite_pat",   "title": "Sécurité patrimoniale",       "icon": "🔐", "step": 3,
     "desc": "Diversification et prévoyance"},
    # Section 4 – Gouvernance
    {"id": "gouvernance",    "title": "Gouvernance de la holding",   "icon": "🏢", "step": 4,
     "desc": "Règles de fonctionnement et de décision"},
    {"id": "continuite",     "title": "Continuité opérationnelle",   "icon": "⚙️", "step": 4,
     "desc": "Dépendance au dirigeant actuel"},
    {"id": "succession_mgmt","title": "Succession managériale",      "icon": "👤", "step": 4,
     "desc": "Préparation et calendrier du successeur"},
    # Section 5 – Fiscal & Suivi
    {"id": "fiscalite",      "title": "Fiscalité",                   "icon": "💶", "step": 5,
     "desc": "Simulation fiscale et Pacte Dutreil"},
    {"id": "dutreil",        "title": "Vigilance Dutreil",           "icon": "⚠️", "step": 5,
     "desc": "Conditions et suivi du Pacte Dutreil"},
    {"id": "suivi",          "title": "Suivi de la stratégie",       "icon": "📅", "step": 5,
     "desc": "Revue annuelle et formalisation"},
    # Synthèse
    {"id": "synthese",       "title": "Synthèse du diagnostic",      "icon": "✨", "step": None,
     "desc": "Récapitulatif et accès aux résultats détaillés"},
]

SECTION_NAMES: Dict[Any, str] = {
    1: "Cadrage & objectifs",
    2: "Famille",
    3: "Patrimoine",
    4: "Gouvernance",
    5: "Fiscal & Suivi",
    None: "Synthèse",
}


def is_subpage_active(sp_id: str, draft: Dict) -> bool:
    # Objectifs : toujours chercher dans draft, puis cache, puis answers validées
    _obj_raw = (draft.get("objectifs") or
                st.session_state.get("_objectifs_cache") or
                st.session_state.answers.get("objectifs") or [])
    objectifs = set(_obj_raw)
    nb_enfants = int(draft.get("nb_enfants") or 0)
    poids = int(draft.get("poids_entreprise") or 0)
    valeur = int(draft.get("valeur_entreprise") or 0)
    if sp_id == "poids":
        return bool(objectifs)
    if sp_id == "repreneur":
        return nb_enfants > 0 and "Transmettre l'entreprise" in objectifs
    if sp_id == "dialogue":
        return nb_enfants >= 2
    if sp_id == "conjoint":
        return draft.get("conjoint_present") == YES
    if sp_id == "succession_civ":
        return nb_enfants >= 2 or draft.get("famille_recomposee") == YES
    if sp_id == "securite_pat":
        show_div = poids >= 40 or "Diversifier le patrimoine" in objectifs
        show_prev = (show_div or draft.get("conjoint_present") == YES
                     or "Protéger le conjoint et les proches" in objectifs
                     or draft.get("besoin_revenus_famille") in ["Moyen", "Élevé"]
                     or draft.get("urgence_evenement") == YES)
        return show_div or show_prev
    if sp_id == "gouvernance":
        return ("Conserver le contrôle familial" in objectifs
                or nb_enfants >= 2
                or draft.get("besoin_revenus_famille") in ["Moyen", "Élevé"]
                or "Préserver l'équité entre les héritiers" in objectifs)
    if sp_id == "succession_mgmt":
        return draft.get("heritier_repreneur") in [YES, UNCERTAIN]
    if sp_id == "fiscalite":
        return "Optimiser la fiscalité" in objectifs or valeur >= 1_000_000
    if sp_id == "dutreil":
        return draft.get("pacte_dutreil") == YES
    return True


def get_active_subpages(draft: Dict) -> List[Dict]:
    return [sp for sp in ALL_SUBPAGES if is_subpage_active(sp["id"], draft)]


def get_subpage_idx(sp_id: str, active: List[Dict]) -> int:
    ids = [p["id"] for p in active]
    return ids.index(sp_id) if sp_id in ids else 0


def _flush_widgets_to_draft() -> None:
    """Copie TOUS les widgets visibles dans draft_answers avant de changer de page."""
    # 1. Restaure d'abord les objectifs depuis le cache si draft les a perdus
    if not st.session_state.draft_answers.get("objectifs"):
        cached = (st.session_state.get("_objectifs_cache") or
                  st.session_state.answers.get("objectifs") or [])
        if cached:
            st.session_state.draft_answers["objectifs"] = list(cached)
    # 2. Copie les widgets standards (sauf objectifs/objective_weights gérés séparément)
    _skip = {"objectifs", "objective_weights"}
    for step_keys in STEP_KEYS.values():
        for key in step_keys:
            if key in _skip:
                continue
            wk = f"w_{key}"
            if wk in st.session_state:
                st.session_state.draft_answers[key] = st.session_state[wk]
    # 3. Persiste les objectifs dans answers et cache
    if st.session_state.draft_answers.get("objectifs"):
        st.session_state.answers["objectifs"] = st.session_state.draft_answers["objectifs"]
        st.session_state["_objectifs_cache"] = st.session_state.draft_answers["objectifs"]
    # 4. Pondérations — lues depuis les radios au moment de naviguer (pas de on_change)
    for obj in (st.session_state.draft_answers.get("objectifs") or []):
        wk = f"w_objective_weight_{safe_key(obj)}"
        if wk in st.session_state:
            st.session_state.draft_answers.setdefault("objective_weights", {})[obj] = st.session_state[wk]


def navigate_next(current_idx: int, active_pages: List[Dict]) -> None:
    _flush_widgets_to_draft()
    current = active_pages[current_idx]

    # Validation inline : impossible de quitter "objectifs" sans avoir sélectionné au moins un objectif
    if current["id"] == "objectifs":
        obj_selected = (
            st.session_state.draft_answers.get("objectifs") or
            st.session_state.get("_objectifs_cache") or
            st.session_state.answers.get("objectifs") or []
        )
        if not obj_selected:
            st.session_state.nav_error = "⚠️ Sélectionne au moins un objectif pour continuer."
            st.rerun()
            return

    is_last = current_idx + 1 >= len(active_pages)
    if not is_last:
        next_pg = active_pages[current_idx + 1]
        crosses = current["step"] is not None and next_pg["step"] != current["step"]
    else:
        crosses = current["step"] is not None
    if crosses:
        ok, msg = save_step(current["step"])
        if not ok:
            st.session_state.nav_error = msg
            if "objectif" in msg.lower():
                st.session_state.current_subpage = "objectifs"
            st.rerun()
            return
    st.session_state.nav_error = None
    st.session_state.current_subpage = "synthese" if is_last else active_pages[current_idx + 1]["id"]
    st.rerun()


def navigate_prev(current_idx: int, active_pages: List[Dict]) -> None:
    _flush_widgets_to_draft()
    if current_idx > 0:
        st.session_state.nav_error = None
        st.session_state.current_subpage = active_pages[current_idx - 1]["id"]
        st.rerun()


# ─── CSS ─────────────────────────────────────────────────────────────────────

st.markdown("""
<style>
.main .block-container { padding-top: 0.6rem; max-width: 980px; }

.hero {
    background: linear-gradient(135deg, #0ea5e9 0%, #06b6d4 45%, #10b981 100%);
    color: white; padding: 18px 26px; border-radius: 16px; margin-bottom: 18px;
    box-shadow: 0 4px 20px rgba(14,165,233,.2);
}
.hero h1 { margin: 0; color: white; font-size: 1.4rem; font-weight: 700; }
.hero p  { margin: 4px 0 0; color: rgba(255,255,255,.9); font-size: .86rem; }

.page-hdr {
    display: flex; align-items: center; gap: 16px;
    background: linear-gradient(135deg, #f0f9ff 0%, #f0fdf4 100%);
    border: 1px solid #bae6fd; border-radius: 16px;
    padding: 16px 22px; margin-bottom: 20px;
}
.page-hdr-icon  { font-size: 2.2rem; line-height: 1; flex-shrink: 0; }
.page-hdr-title { font-size: 1.1rem; font-weight: 700; color: #0c4a6e; }
.page-hdr-desc  { font-size: .82rem; color: #0369a1; margin-top: 3px; }
.page-hdr-badge {
    font-size: .68rem; font-weight: 700; background: #0ea5e9; color: white !important;
    padding: 2px 8px; border-radius: 999px; margin-left: 8px; vertical-align: middle;
}

.sc-title {
    font-size: .72rem; font-weight: 700; text-transform: uppercase; letter-spacing: .06em;
    color: #0369a1; background: #f0f9ff; border-left: 3px solid #38bdf8;
    padding: 7px 12px; border-radius: 0 6px 6px 0; margin: 18px 0 12px 0;
}

.card {
    background: white !important; border-radius: 14px; padding: 18px 20px; margin: 10px 0;
    box-shadow: 0 2px 8px rgba(14,165,233,.08); border: 1px solid #e0f2fe;
}
.card h3, .card p, .card li, .card div, .card strong { color: #0c172a !important; }
.card h3 { margin-top: 0; margin-bottom: 8px; }

.badge {
    color: white !important; padding: 4px 10px; border-radius: 999px;
    font-weight: 700; font-size: .75rem; display: inline-block;
}

.importance-card {
    background: #fafffe; border: 1px solid #a7f3d0; border-radius: 12px;
    padding: 14px 16px; margin: 8px 0;
}
.importance-help {
    background: #f0f9ff; border: 1px solid #bae6fd; border-radius: 8px;
    padding: 9px 12px; color: #0369a1; margin: 6px 0 14px 0; font-size: .84rem;
}

.justification-table { width: 100%; border-collapse: collapse; font-size: .85rem; }
.justification-table th {
    background: #f0f9ff; color: #0c4a6e; text-align: left;
    padding: 8px 12px; border: 1px solid #e0f2fe;
}
.justification-table td { padding: 8px 12px; border: 1px solid #e0f2fe; vertical-align: top; }

.solution-grid { display: grid; grid-template-columns: repeat(3,1fr); gap: 14px; margin-top: 12px; }
@media(max-width:1000px) { .solution-grid { grid-template-columns: 1fr; } }

.small-note { color: #475569 !important; font-size: .87rem; }
</style>
""", unsafe_allow_html=True)

# ─── Hero ────────────────────────────────────────────────────────────────────

st.markdown("""
<div class="hero">
  <h1>🏛️ Diagnostic Holding Familiale</h1>
  <p>Système expert interactif — Objectif → Risque → Outil → Prévention → Suivi</p>
</div>
""", unsafe_allow_html=True)

# ─── Calcul des risques (réponses validées uniquement) ───────────────────────

validated_answers = st.session_state.answers
df, evidence = calculate_risks(validated_answers)
detected_df = df[df["Score"] > 0].copy()
zero_df = df[df["Score"] == 0].copy()

# ─── Helpers rendu ───────────────────────────────────────────────────────────

def section_card(title: str, icon: str = "") -> None:
    prefix = f"{icon} " if icon else ""
    st.markdown(f'<div class="sc-title">{prefix}{title}</div>', unsafe_allow_html=True)


def render_view_navigation(current_view: str) -> None:
    idx = VIEWS.index(current_view)
    st.divider()
    c1, _, c3 = st.columns([1.8, 5, 1.8])
    with c1:
        if idx > 0:
            if st.button(f"← {VIEWS[idx - 1]}", key=f"prev_view_{current_view}"):
                st.session_state.app_page = VIEWS[idx - 1]
                st.rerun()
    with c3:
        if idx < len(VIEWS) - 1:
            if st.button(f"{VIEWS[idx + 1]} →", key=f"next_view_{current_view}", type="primary"):
                st.session_state.app_page = VIEWS[idx + 1]
                st.rerun()


def render_page_header(sp: Dict) -> None:
    step = sp.get("step")
    section = SECTION_NAMES.get(step, "") if step is not None else ""
    badge = (f'<span class="page-hdr-badge">{section}</span>' if section else "")
    st.markdown(
        f'<div class="page-hdr">'
        f'<div class="page-hdr-icon">{sp["icon"]}</div>'
        f'<div>'
        f'<div class="page-hdr-title">{sp["title"]}{badge}</div>'
        f'<div class="page-hdr-desc">{sp["desc"]}</div>'
        f'</div>'
        f'</div>',
        unsafe_allow_html=True,
    )


def render_subpage(sp_id: str) -> None:
    """Affiche les questions de la sous-page courante."""

    if sp_id == "dossier":
        c1, c2, c3 = st.columns(3)
        with c1:
            text_input_field("Nom du dossier / client", "client_name", "Ex. Famille Martin")
            number_input("Âge du dirigeant", "client_age", min_value=0, max_value=100, step=1)
        with c2:
            text_input_field("Nom de l'entreprise", "company_name", "Ex. Martin Industrie")
            selectbox("Forme juridique", "company_form",
                      [UNKNOWN, "SAS", "SARL", "SA", "SNC", "Entreprise individuelle", "Autre"])
        with c3:
            text_input_field("Activité / secteur", "company_activity", "Ex. industrie, bâtiment…")
            text_input_field("CGP / cabinet", "cgp_name", "Ex. Cabinet Dupont")
        c4, c5 = st.columns(2)
        with c4:
            text_input_field("Date de l'entretien", "entretien_date", "Ex. 21/06/2026")
        with c5:
            selectbox("Qualité des informations recueillies", "qualite_information",
                      [UNKNOWN, "Confirmées", "Partiellement confirmées", "À vérifier"])

    elif sp_id == "maturite":
        c1, c2 = st.columns(2)
        with c1:
            selectbox("Niveau de maturité du projet", "maturite_projet",
                      [UNKNOWN, "Simple réflexion", "Projet envisagé",
                       "Transmission prévue", "Transmission urgente"])
            selectbox("Horizon de transmission envisagé", "delai_transmission",
                      [UNKNOWN, "Moins de 12 mois", "1 à 3 ans", "Plus de 3 ans", "Pas défini"])
        with c2:
            yes_no_unknown("Événement d'urgence ou de fragilité ?", "urgence_evenement")
            st.caption(
                "Maladie, accident, conflit imminent, pression extérieure… "
                "tout signal qui accélère la nécessité d'agir."
            )

    elif sp_id == "objectifs":
        sync_widget_from_draft("objectifs")

        def _on_objectifs_change():
            val = st.session_state.get("w_objectifs") or []
            st.session_state.draft_answers["objectifs"] = val
            # Stocke dans une clé persistante non-widget (résiste aux reruns)
            st.session_state["_objectifs_cache"] = val

        st.multiselect(
            "Quels objectifs le dirigeant poursuit-il principalement ?",
            OBJECTIVE_DISPLAY_ORDER,
            key="w_objectifs",
            placeholder="Sélectionner un ou plusieurs objectifs",
            on_change=_on_objectifs_change,
        )
        selected = list(get_draft("objectifs") or [])
        # Alimente le cache à chaque render (même sans on_change)
        if selected:
            st.session_state["_objectifs_cache"] = selected
            st.session_state.draft_answers["objectifs"] = selected
        if not selected:
            st.info("⚠️ Sélectionne au moins un objectif pour activer le scoring des risques.")
        else:
            st.success(
                f"✅ {len(selected)} objectif(s) sélectionné(s). "
                "Une page de pondération s'affichera à l'étape suivante."
            )

    elif sp_id == "poids":
        # Priorité : cache persistant → draft → answers (w_objectifs non rendu ici)
        selected = list(
            st.session_state.get("_objectifs_cache") or
            st.session_state.draft_answers.get("objectifs") or
            st.session_state.answers.get("objectifs") or []
        )
        # Restaure le draft si nécessaire
        if selected and not st.session_state.draft_answers.get("objectifs"):
            st.session_state.draft_answers["objectifs"] = selected
        if not selected:
            st.warning("Aucun objectif sélectionné. Reviens à la page précédente.")
        else:
            st.markdown(
                '<div class="importance-help">'
                '<strong>Important</strong> : pondération normale &nbsp;·&nbsp; '
                '<strong>Très important</strong> : renforcement modéré &nbsp;·&nbsp; '
                '<strong>Prioritaire</strong> : renforcement fort des risques associés.'
                '</div>',
                unsafe_allow_html=True,
            )
            for objective in selected:
                st.markdown('<div class="importance-card">', unsafe_allow_html=True)
                objective_weight_buttons(objective)
                st.markdown('</div>', unsafe_allow_html=True)

    elif sp_id == "rapport":
        c1, c2 = st.columns(2)
        with c1:
            selectbox("Orientation du rapport exporté", "rapport_orientation",
                      ["Équilibré", "Très pédagogique",
                       "Synthétique et décisionnel", "Approfondi et technique"])
            st.caption("Modifie la structure et le niveau de détail du rapport Word.")
        with c2:
            selectbox("Niveau de détail", "niveau_detail",
                      ["Synthétique", "Détaillé", "Très détaillé"])
        text_area_field(
            "Objectif exprimé par le client (ses propres mots)", "objectif_libre",
            "Ex. transmettre progressivement à mon fils sans léser mes autres enfants"
        )
        c1, c2 = st.columns(2)
        with c1:
            text_area_field("Attentes particulières du client", "attentes_client",
                            "Ex. éviter les conflits, protéger le conjoint…")
        with c2:
            text_area_field("Contraintes ou préférences", "contraintes_client",
                            "Ex. refus d'ouvrir le capital, souhait de simplicité…")
        c1, c2 = st.columns(2)
        with c1:
            text_area_field("Personnes à associer à la réflexion", "personnes_a_associer",
                            "Ex. conjoint, enfants, notaire, avocat…")
        with c2:
            text_area_field("Observations du CGP", "observations",
                            "Notes de contexte utiles pour le rapport", height=90)

    elif sp_id == "famille":
        c1, c2, c3 = st.columns(3)
        with c1:
            number_input("Nombre d'enfants", "nb_enfants", min_value=0, max_value=12, step=1)
        with c2:
            yes_no_unknown("Conjoint présent ?", "conjoint_present")
        with c3:
            yes_no_unknown("Famille recomposée ?", "famille_recomposee")
        if get_draft("conjoint_present") == YES:
            st.markdown("<br>", unsafe_allow_html=True)
            yes_no_unknown(
                "Le conjoint adhère-t-il au projet de transmission ?",
                "accord_conjoint", [UNKNOWN, YES, NO, UNCERTAIN]
            )
        nb = int(get_draft("nb_enfants") or 0)
        obj_d = set(get_draft("objectifs") or [])
        hints = []
        if nb > 0 and "Transmettre l'entreprise" in obj_d:
            hints.append("📌 Repreneur & transmission")
        if nb >= 2:
            hints.append("📌 Dialogue familial")
        if get_draft("conjoint_present") == YES:
            hints.append("📌 Protection du conjoint")
        if nb >= 2 or get_draft("famille_recomposee") == YES:
            hints.append("📌 Sécurisation successorale")
        if hints:
            st.caption("Pages qui s'ajouteront automatiquement : " + " · ".join(hints))

    elif sp_id == "repreneur":
        yes_no_unknown("Un héritier repreneur est-il identifié ?", "heritier_repreneur",
                       [UNKNOWN, YES, NO, UNCERTAIN])
        if get_draft("heritier_repreneur") == YES and int(get_draft("nb_enfants") or 0) >= 2:
            st.divider()
            yes_no_unknown(
                "Les autres héritiers sont-ils aussi impliqués dans l'entreprise ?",
                "autres_heritiers_actifs"
            )
            selectbox("Volonté probable des héritiers non repreneurs", "volonte_non_repreneurs",
                      [UNKNOWN, "Rester associés", "Sortir du capital",
                       "Recevoir principalement une compensation", "Incertain / non abordé"])
            yes_no_unknown("Une soulte (compensation financière) est-elle envisagée ?",
                           "soulte_envisagee")
            if get_draft("soulte_envisagee") == YES:
                yes_no_unknown(
                    "La capacité de financement de la soulte est-elle validée ?",
                    "capacite_financement_soulte"
                )

    elif sp_id == "dialogue":
        yes_no_unknown(
            "Un dialogue familial a-t-il déjà été organisé entre les héritiers ?",
            "dialogue_familial"
        )
        st.caption("Un dialogue précoce réduit fortement les risques de conflit successoral.")

    elif sp_id == "conjoint":
        st.caption("Ces questions sont prises en compte dès qu'un conjoint est présent.")
        c1, c2, c3 = st.columns(3)
        with c1:
            yes_no_unknown("Le conjoint dépend-il financièrement du dirigeant ?",
                           "conjoint_dependant")
        with c2:
            yes_no_unknown("Un dispositif de protection du conjoint est-il prévu ?",
                           "protection_conjoint_prevue")
        with c3:
            yes_no_unknown("Le régime matrimonial a-t-il été analysé ou adapté ?",
                           "regime_matrimonial_adapte")

    elif sp_id == "succession_civ":
        c1, c2 = st.columns(2)
        with c1:
            yes_no_unknown("Une valorisation indépendante des titres est-elle prévue ?",
                           "valorisation_independante")
        with c2:
            yes_no_unknown(
                "Un audit civil et successoral a-t-il été réalisé avec un notaire ?",
                "audit_civil"
            )

    elif sp_id == "patrimoine":
        c1, c2, c3 = st.columns(3)
        with c1:
            number_input("Valeur estimée de l'entreprise (€)", "valeur_entreprise",
                         min_value=0, step=100_000)
        with c2:
            slider("Poids de l'entreprise dans le patrimoine (%)", "poids_entreprise", 0, 100)
        with c3:
            selectbox("Actifs liquides hors entreprise", "actifs_liquides",
                      [UNKNOWN, "Faible", "Moyen", "Élevé"])
        c4, c5 = st.columns(2)
        with c4:
            selectbox("Endettement personnel ou familial", "endettement_familial",
                      [UNKNOWN, "Faible", "Moyen", "Élevé"])
        with c5:
            selectbox("Besoin de revenus réguliers pour la famille", "besoin_revenus_famille",
                      [UNKNOWN, "Faible", "Moyen", "Élevé"])

    elif sp_id == "securite_pat":
        c1, c2 = st.columns(2)
        with c1:
            yes_no_unknown("Une diversification patrimoniale est-elle déjà organisée ?",
                           "diversification")
        with c2:
            yes_no_unknown("Une prévoyance ou assurance décès est-elle prévue ?", "prevoyance")

    elif sp_id == "gouvernance":
        c1, c2 = st.columns(2)
        with c1:
            yes_no_unknown("Des clauses d'entrée/sortie des titres sont-elles prévues ?",
                           "clauses_entree_sortie")
            yes_no_unknown("La gouvernance de la holding est-elle formalisée ?",
                           "gouvernance_formalisee")
        with c2:
            yes_no_unknown("La holding réunira-t-elle des associés actifs et non actifs ?",
                           "associes_actifs_passifs")
            yes_no_unknown("Une politique de distribution de dividendes est-elle définie ?",
                           "politique_dividendes_definie")

    elif sp_id == "continuite":
        yes_no_unknown(
            "L'entreprise est-elle fortement dépendante du dirigeant actuel ?",
            "entreprise_dependante_dirigeant"
        )

    elif sp_id == "succession_mgmt":
        c1, c2 = st.columns(2)
        with c1:
            yes_no_unknown("Le successeur est-il préparé à reprendre la direction ?",
                           "successeur_prepare")
        with c2:
            yes_no_unknown("Un calendrier de transmission du pouvoir est-il prévu ?",
                           "calendrier_transmission")

    elif sp_id == "fiscalite":
        c1, c2 = st.columns(2)
        with c1:
            yes_no_unknown("Une simulation fiscale a-t-elle été réalisée ou programmée ?",
                           "simulation_fiscale")
        with c2:
            yes_no_unknown("Un Pacte Dutreil est-il envisagé ou déjà mis en place ?",
                           "pacte_dutreil")
        if get_draft("pacte_dutreil") == YES:
            st.info("📌 Une page « Vigilance Dutreil » va s'afficher automatiquement.")

    elif sp_id == "dutreil":
        st.caption(
            "Le Pacte Dutreil est un dispositif sensible aux conditions strictes. "
            "Ces questions identifient les points de vigilance."
        )
        c1, c2, c3 = st.columns(3)
        with c1:
            yes_no_unknown("La qualification de holding animatrice est-elle établie ?",
                           "holding_animatrice", [UNKNOWN, YES, NO, UNCERTAIN])
        with c2:
            yes_no_unknown("Un audit Dutreil a-t-il été réalisé ou programmé ?", "audit_dutreil")
        with c3:
            yes_no_unknown("Un suivi des engagements de conservation est-il prévu ?",
                           "suivi_engagements")

    elif sp_id == "suivi":
        c1, c2 = st.columns(2)
        with c1:
            yes_no_unknown("Une revue annuelle de la stratégie patrimoniale est-elle prévue ?",
                           "suivi_annuel")
        with c2:
            yes_no_unknown(
                "Un rapport écrit de diagnostic et de recommandation est-il prévu ?",
                "formalisation_rapport"
            )

    elif sp_id == "synthese":
        n_valid = len(st.session_state.validated_steps.intersection({1, 2, 3, 4, 5}))
        if n_valid == 0:
            st.warning(
                "⚠️ Aucune section validée. Complète et valide les sections "
                "du questionnaire pour voir le diagnostic."
            )
        else:
            gaps = missing_data(validated_answers)
            if gaps:
                with st.expander(
                    f"⚠️ {len(gaps)} information(s) à compléter pour affiner le diagnostic"
                ):
                    for g in gaps:
                        st.write(f"- {g}")
            if detected_df.empty:
                st.success("✅ Aucun risque détecté sur la base des réponses validées.")
            else:
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("Risques détectés", len(detected_df))
                c2.metric("Critiques", int((detected_df["Niveau"] == "Critique").sum()))
                c3.metric("Élevés", int((detected_df["Niveau"] == "Élevé").sum()))
                c4.metric("Score max", int(detected_df["Score"].max()))
                st.dataframe(
                    detected_df[["Objectif", "Risque", "Score", "Niveau"]],
                    use_container_width=True, hide_index=True,
                )


def render_nav_buttons(current_idx: int, active_pages: List[Dict]) -> None:
    current = active_pages[current_idx]
    is_last = current_idx + 1 >= len(active_pages)

    if current["id"] == "synthese":
        if st.button("Voir les résultats complets →", type="primary",
                     key="synth_to_results", use_container_width=True):
            st.session_state.app_page = "Résultats et solutions"
            st.rerun()
        return

    if not is_last:
        next_pg = active_pages[current_idx + 1]
        crosses = current["step"] is not None and next_pg["step"] != current["step"]
    else:
        crosses = current["step"] is not None

    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    c1, _, c3 = st.columns([1.5, 4, 1.5])
    with c1:
        if current_idx > 0:
            if st.button("← Précédent", key=f"prev_{current['id']}"):
                navigate_prev(current_idx, active_pages)
    with c3:
        if is_last:
            btn_label = "Terminer →"
        elif crosses:
            btn_label = "Valider & continuer →"
        else:
            btn_label = "Suivant →"
        if st.button(btn_label, type="primary", key=f"next_{current['id']}"):
            navigate_next(current_idx, active_pages)


# ─── Sidebar ─────────────────────────────────────────────────────────────────

page = st.session_state.get("app_page", VIEWS[0])

with st.sidebar:
    st.markdown("### Navigation")
    _btn_icons = {
        "Questionnaire adaptatif": "📋",
        "Résultats et solutions": "📊",
        "Exporter": "⬇️",
    }
    for _v in VIEWS[:3]:
        _active = _v == page
        _lbl = (_btn_icons.get(_v, "") + " " + _v).strip()
        if st.button(_lbl, key=f"nav_{safe_key(_v)}", use_container_width=True,
                     type="primary" if _active else "secondary"):
            st.session_state.app_page = _v
            st.rerun()

    st.divider()

    _draft = st.session_state.draft_answers
    _active_pages = get_active_subpages(_draft)
    _current_sp = st.session_state.get("current_subpage", "dossier")
    if _current_sp not in [p["id"] for p in _active_pages]:
        _current_sp = _active_pages[0]["id"] if _active_pages else "dossier"
    _current_idx = get_subpage_idx(_current_sp, _active_pages)

    if page == "Questionnaire adaptatif":
        _pct = _current_idx / max(len(_active_pages) - 1, 1) * 100
        st.markdown(
            f'<div style="font-size:.76rem;color:#0369a1;font-weight:600;margin-bottom:4px">'
            f'Page {_current_idx + 1} / {len(_active_pages)}</div>'
            f'<div style="background:#bae6fd;border-radius:999px;height:6px;margin-bottom:10px">'
            f'<div style="background:linear-gradient(90deg,#38bdf8,#22c55e);border-radius:999px;'
            f'height:6px;width:{_pct:.0f}%"></div></div>',
            unsafe_allow_html=True,
        )
        _prev_step: Any = object()
        for _i, _sp in enumerate(_active_pages):
            if _sp["step"] != _prev_step:
                _prev_step = _sp["step"]
                _sname = SECTION_NAMES.get(_sp["step"], "Synthèse")
                _step_done = _sp["step"] in st.session_state.validated_steps if _sp["step"] else False
                st.markdown(
                    f"<div style='font-size:.68rem;font-weight:700;color:#0369a1;"
                    f"text-transform:uppercase;letter-spacing:.05em;margin:8px 0 2px 0'>"
                    f"{'✅' if _step_done else '◦'} {_sname}</div>",
                    unsafe_allow_html=True,
                )
            _is_cur = _sp["id"] == _current_sp
            _is_done = _i < _current_idx
            _dot = "▶" if _is_cur else ("✓" if _is_done else "·")
            _color = "#0ea5e9" if _is_cur else ("#10b981" if _is_done else "#9ca3af")
            _wt = "700" if _is_cur else "400"
            st.markdown(
                f"<div style='font-size:.77rem;color:{_color};font-weight:{_wt};"
                f"padding:1px 0 1px 10px;line-height:1.65'>{_dot} {_sp['title']}</div>",
                unsafe_allow_html=True,
            )
    else:
        _n_done = len(st.session_state.validated_steps.intersection({1, 2, 3, 4, 5}))
        st.markdown(
            f'<div style="font-size:.76rem;color:#0369a1;font-weight:600;margin-bottom:4px">'
            f'{_n_done} / 5 sections validées</div>'
            f'<div style="background:#bae6fd;border-radius:999px;height:6px;margin-bottom:10px">'
            f'<div style="background:linear-gradient(90deg,#38bdf8,#22c55e);border-radius:999px;'
            f'height:6px;width:{int(_n_done/5*100)}%"></div></div>',
            unsafe_allow_html=True,
        )

    st.divider()
    for _v in VIEWS[3:]:
        _active = _v == page
        if st.button(_v, key=f"nav_{safe_key(_v)}", use_container_width=True,
                     type="primary" if _active else "secondary"):
            st.session_state.app_page = _v
            st.rerun()

    st.divider()
    if st.button("🔄 Réinitialiser le diagnostic", type="secondary", use_container_width=True):
        reset_app()
        st.session_state.setdefault("current_subpage", "dossier")
        st.rerun()
    st.caption("Les réponses ne modifient pas les résultats tant qu'elles ne sont pas validées.")

# ═══════════════════════════════════════════════════════════════════════════════
# PAGES
# ═══════════════════════════════════════════════════════════════════════════════

if page == "Questionnaire adaptatif":
    draft = st.session_state.draft_answers
    active_pages = get_active_subpages(draft)

    current_sp = st.session_state.get("current_subpage", "dossier")
    active_ids = [p["id"] for p in active_pages]
    if current_sp not in active_ids:
        current_sp = active_ids[0] if active_ids else "dossier"
        st.session_state.current_subpage = current_sp

    current_idx = get_subpage_idx(current_sp, active_pages)
    current_page_def = active_pages[current_idx]

    if st.session_state.get("nav_error"):
        st.error(st.session_state.nav_error)

    pct_main = current_idx / max(len(active_pages) - 1, 1) * 100
    st.markdown(
        f'<div style="font-size:.75rem;color:#0369a1;font-weight:600;margin-bottom:4px">'
        f'Page {current_idx+1} / {len(active_pages)}</div>'
        f'<div style="background:#bae6fd;border-radius:999px;height:5px;margin-bottom:14px">'
        f'<div style="background:linear-gradient(90deg,#38bdf8,#22c55e);border-radius:999px;'
        f'height:5px;width:{pct_main:.0f}%"></div></div>',
        unsafe_allow_html=True,
    )

    render_page_header(current_page_def)
    render_subpage(current_sp)
    render_nav_buttons(current_idx, active_pages)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE : Résultats et solutions
# ═══════════════════════════════════════════════════════════════════════════════

elif page == "Résultats et solutions":
    st.subheader("Résultats et solutions proposées")
    st.caption("Calculés uniquement à partir des réponses validées.")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Risques détectés", len(detected_df))
    c2.metric("Inexistants", len(zero_df))
    c3.metric("Critiques",
              int((detected_df["Niveau"] == "Critique").sum()) if not detected_df.empty else 0)
    c4.metric("Élevés",
              int((detected_df["Niveau"] == "Élevé").sum()) if not detected_df.empty else 0)

    gaps = missing_data(validated_answers)
    if gaps:
        with st.expander(f"⚠️ {len(gaps)} information(s) manquante(s) pour affiner le diagnostic"):
            for g in gaps:
                st.write(f"- {g}")

    if not detected_df.empty:
        section_card("Visualisation des scores", "📊")
        chart_df = detected_df[["Risque", "Score"]].sort_values("Score", ascending=True)
        st.bar_chart(chart_df, x="Risque", y="Score", horizontal=True, color="#0ea5e9")

    section_card("Tableau des risques", "📋")
    st.dataframe(
        df[["Objectif", "Risque", "Probabilité", "Gravité", "Score",
            "Niveau", "Pondération objectif", "Signaux retenus"]],
        use_container_width=True, hide_index=True,
        column_config={"Signaux retenus": st.column_config.TextColumn(width="large")},
    )

    section_card("Solutions et outils par risque détecté", "🛠️")
    if detected_df.empty:
        st.info("Aucun risque détecté. Complète le questionnaire et valide les sections.")
    else:
        for _, row in detected_df.iterrows():
            definition = RISK_DEFINITIONS[row["Code"]]
            color = PRIORITY_COLORS[row["Niveau"]]
            body = f"""
            <p>{priority_badge(row['Niveau'])} &nbsp;
               <strong>Score :</strong> {int(row['Score'])} &nbsp;
               <strong>Probabilité :</strong> {score_to_label(int(row['Probabilité']))} &nbsp;
               <strong>Gravité :</strong> {score_to_label(int(row['Gravité']))}</p>
            <p><strong>Objectif concerné :</strong> {definition.objectif}</p>
            <p><strong>Pourquoi ce risque est activé :</strong></p>
            {render_list(evidence.get(row['Code'], []))}
            <div class="solution-grid">
              <div><strong>Conséquences possibles</strong>{render_list(definition.consequences)}</div>
              <div><strong>Outils à étudier</strong>{render_list(definition.outils)}</div>
              <div><strong>Actions préventives</strong>{render_list(definition.actions_preventives)}</div>
            </div>
            <p><strong>Justification des outils :</strong></p>
            {render_tool_justifications(row['Code'])}
            <p><strong>Professionnels à mobiliser :</strong> {', '.join(definition.professionnels)}</p>
            """
            card(definition.libelle, body, accent=color)

    section_card("Plan d'action hiérarchisé", "🗓️")
    plan_df = build_action_plan(df)
    if plan_df.empty:
        st.info("Le plan d'action apparaît dès que des risques sont détectés.")
    else:
        for horizon in ["Actions immédiates", "Actions à moyen terme", "Suivi annuel"]:
            hdf = plan_df[plan_df["Horizon"] == horizon]
            if not hdf.empty:
                with st.expander(horizon, expanded=(horizon == "Actions immédiates")):
                    st.dataframe(
                        hdf[["Priorité", "Action", "Risques concernés", "Professionnels"]],
                        use_container_width=True, hide_index=True,
                        column_config={
                            "Action": st.column_config.TextColumn(width="large"),
                            "Professionnels": st.column_config.TextColumn(width="medium"),
                        },
                    )

    with st.expander("Risques non détectés (au vu des réponses validées)"):
        st.dataframe(zero_df[["Objectif", "Risque", "Niveau"]],
                     use_container_width=True, hide_index=True)

    render_view_navigation("Résultats et solutions")

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE : Exporter
# ═══════════════════════════════════════════════════════════════════════════════

elif page == "Exporter":
    st.subheader("Exporter le diagnostic")

    n_valid = len(st.session_state.validated_steps.intersection({1, 2, 3, 4, 5}))
    if n_valid == 0:
        st.warning("⚠️ Aucune section validée — le rapport sera vide. Complète d'abord le questionnaire.")
    elif n_valid < 3:
        st.info(f"ℹ️ {n_valid}/5 sections validées. Tu peux exporter, mais le diagnostic sera partiel.")

    section_card("Télécharger", "⬇️")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("**📄 Rapport Word**")
        st.caption("Rapport structuré adapté à l'orientation choisie.")
        try:
            docx_data = create_docx_report(
                validated_answers.get("client_name") or "", validated_answers, df, evidence
            )
            st.download_button(
                "Télécharger le .docx", data=docx_data,
                file_name="rapport_diagnostic_holding_familiale.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary", use_container_width=True,
            )
        except Exception as exc:
            st.error(f"Indisponible : {exc}")
    with c2:
        st.markdown("**📊 Matrice Excel**")
        st.caption("Deux onglets : matrice des risques + plan d'action.")
        try:
            xlsx_data = create_xlsx_matrix(df)
            st.download_button(
                "Télécharger le .xlsx", data=xlsx_data,
                file_name="matrice_risques_holding_familiale.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
        except Exception as exc:
            st.error(f"Indisponible : {exc}")
    with c3:
        st.markdown("**📝 Rapport Markdown**")
        st.caption("Version texte brut compatible Obsidian, Notion ou VS Code.")
        try:
            md_text = markdown_report(
                validated_answers.get("client_name") or "", validated_answers, df
            )
            st.download_button(
                "Télécharger le .md", data=md_text.encode("utf-8"),
                file_name="rapport_diagnostic_holding_familiale.md",
                mime="text/markdown", use_container_width=True,
            )
        except Exception as exc:
            st.error(f"Indisponible : {exc}")

    st.divider()
    orientation = validated_answers.get("rapport_orientation", "Équilibré")
    section_card(f"Contenu du rapport — Orientation : {orientation}", "ℹ️")
    st.info(
        "**Synthétique** : risques prioritaires et décisions uniquement.  \n"
        "**Pédagogique** : explications accessibles, lexique, pédagogie famille.  \n"
        "**Équilibré** : analyse complète avec scoring et plan d'action.  \n"
        "**Technique** : scoring détaillé, signaux retenus, validations professionnelles."
    )

    if not detected_df.empty:
        section_card("Aperçu — Risques détectés", "📋")
        st.dataframe(
            detected_df[["Objectif", "Risque", "Score", "Niveau",
                          "Outils à étudier", "Actions préventives"]],
            use_container_width=True, hide_index=True,
            column_config={
                "Outils à étudier": st.column_config.TextColumn(width="large"),
                "Actions préventives": st.column_config.TextColumn(width="large"),
            },
        )

    plan_preview_df = build_action_plan(df)
    if not plan_preview_df.empty:
        section_card("Aperçu — Plan d'action", "🗓️")
        st.dataframe(
            plan_preview_df, use_container_width=True, hide_index=True,
            column_config={"Action": st.column_config.TextColumn(width="large")},
        )

    render_view_navigation("Exporter")

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE : Règles de décision
# ═══════════════════════════════════════════════════════════════════════════════

elif page == "Règles de décision":
    st.subheader("Règles de décision du système expert")
    st.write("Chaque réponse validée ajoute des points à certains risques. Aucun signal = niveau Inexistant.")
    st.code("""
SI une section n'est pas validée
ALORS ses réponses ne sont pas utilisées dans le scoring.

SI aucun objectif n'est validé
ALORS tous les risques restent à 0.

SI un objectif est pondéré « Très important » ou « Prioritaire »
ALORS les risques associés sont renforcés dans le scoring.

SI le contrôle familial est recherché ET plusieurs héritiers existent
ALORS le risque de dilution augmente.

SI un enfant reprend ET les autres héritiers ne sont pas impliqués
ALORS le risque de conflit repreneur / non repreneurs augmente.

SI une soulte est envisagée MAIS son financement n'est pas validé
ALORS le risque de liquidité augmente fortement.

SI un Pacte Dutreil est envisagé ET la holding animatrice est incertaine
ALORS le risque de remise en cause du Dutreil augmente fortement.

SI le conjoint dépend financièrement du dirigeant ET aucune protection n'est prévue
ALORS le risque de fragilisation du conjoint augmente.

SI nb_enfants = 0 ET objectif = Transmettre l'entreprise
ALORS le risque successeur augmente (aucun successeur identifiable).

SI famille recomposée
ALORS risques de conflit héritiers, contestation et conjoint augmentent.

SI valeur entreprise >= 5 M€
ALORS enjeux fiscaux, Dutreil et contestation activés même hors objectif déclaré.
    """.strip(), language="text")
    rules_export = {
        code: {
            "objectif": d.objectif,
            "risque": d.libelle,
            "gravite_base": d.gravite_base,
            "outils": d.outils,
            "justification_des_outils": get_tool_justifications(code),
            "actions_preventives": d.actions_preventives,
            "professionnels": d.professionnels,
        }
        for code, d in RISK_DEFINITIONS.items()
    }
    st.download_button(
        "Télécharger le dictionnaire des risques (JSON)",
        data=json.dumps(rules_export, ensure_ascii=False, indent=2).encode("utf-8"),
        file_name="dictionnaire_risques_holding_familiale.json",
        mime="application/json",
    )
    st.json(rules_export)
    render_view_navigation("Règles de décision")

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE : Debug
# ═══════════════════════════════════════════════════════════════════════════════

elif page == "Debug validation":
    st.subheader("Debug — Vérification des réponses")
    st.caption("Vue technique : distingue les réponses en cours et les réponses validées utilisées pour le scoring.")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("### Réponses en cours de saisie")
        st.json(st.session_state.get("draft_answers", {}))
    with c2:
        st.markdown("### Réponses validées (scoring)")
        st.json(validated_answers)
    st.markdown("### Sections validées")
    st.write(sorted(list(st.session_state.validated_steps)))
    st.markdown("### Scores calculés")
    st.dataframe(
        df[["Risque", "Probabilité", "Gravité", "Score", "Niveau", "Signaux retenus"]],
        use_container_width=True, hide_index=True,
        column_config={"Signaux retenus": st.column_config.TextColumn(width="large")},
    )
    render_view_navigation("Debug validation")
